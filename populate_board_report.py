from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_board_report(filepath="smartdoc_agent/examples/board_report_v1.docx"):
    doc = Document()

    # Title - Consistent
    title = doc.add_heading("Q3 Financial Performance Review", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Calibri Light"
        run.font.size = Pt(24)
        run.bold = True

    doc.add_paragraph() # Spacer

    # Section 1: Introduction - Inconsistent Heading
    intro_heading = doc.add_heading("1. Introduction & Executive Summary", level=1)
    for run in intro_heading.runs: # Make it different
        run.font.name = "Arial"
        run.font.size = Pt(18)
        run.bold = True

    p1 = doc.add_paragraph("This report outlines the company's financial performance for the third quarter of FY2024. ")
    p1.add_run("Key highlights include a 15% revenue growth year-over-year and successful launch of Project Titan. ")
    # Inconsistent font in a run
    run_p1_inconsistent = p1.add_run("However, operational costs saw an unexpected increase of 5%.")
    run_p1_inconsistent.font.name = "Times New Roman"
    run_p1_inconsistent.font.size = Pt(11)
    p1.paragraph_format.space_after = Pt(6)


    # Section 2: Detailed Financials - Another Inconsistent Heading
    fin_heading = doc.add_heading("SECTION 2: DETAILED FINANCIAL ANALYSIS", level=1) # Different style
    for run in fin_heading.runs:
        run.font.name = "Georgia"
        run.font.size = Pt(16)
        run.italic = True

    p2 = doc.add_paragraph(
        "The detailed financial statements reveal strong performance in the North American market, "
        "offset slightly by challenges in the APAC region. Gross margin remained stable at 45%."
    )
    # Inconsistent paragraph font & spacing
    for run in p2.runs:
        run.font.name = "Verdana"
        run.font.size = Pt(10.5)
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after = Pt(3)
    p2.paragraph_format.line_spacing = 1.25


    # Sub-section with bullet points - inconsistent indentation/style
    doc.add_heading("2.1 Key Metrics", level=2) # Default style for H2 for now
    doc.add_paragraph("Revenue Streams:", style='ListBullet') # Default bullet
    doc.add_paragraph("Subscription Services: $5.2M", style='ListBullet')

    # Manual bullet with different indent/font
    p_manual_bullet = doc.add_paragraph()
    p_manual_bullet.add_run("• Product Sales: $3.8M").font.name = "Courier New"
    p_manual_bullet.paragraph_format.left_indent = Inches(0.75) # Different indent

    doc.add_paragraph("Consulting: $1.1M", style='ListBullet')


    # Section 3: Departmental Review - Table
    doc.add_heading("3. Departmental Review & Variances", level=1) # Default H1
    p3 = doc.add_paragraph(
        "A review of departmental budgets against actuals shows variances primarily in Marketing (overspend) "
        "and R&D (underspend due to hiring delays)."
    )
    for run in p3.runs: # Consistent body text here
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    p3.paragraph_format.space_after = Pt(12)

    # Simple table with no defined style / inconsistent internal formatting
    doc.add_paragraph("Variance Summary:")
    table = doc.add_table(rows=3, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Department'
    hdr_cells[1].text = 'Budgeted'
    hdr_cells[2].text = 'Actual'
    # Inconsistent header font
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.bold = True

    # Data with mixed fonts
    table.cell(1, 0).text = 'Marketing'
    table.cell(1, 1).text = '$500,000'
    table.cell(1, 2).text = '$550,000'
    table.cell(1,2).paragraphs[0].runs[0].font.name="Arial" # one cell different

    table.cell(2, 0).text = 'R&D'
    table.cell(2, 1).text = '$750,000'
    table.cell(2, 2).text = '$700,000'


    # Section 4: Outlook - inconsistent heading and paragraph
    outlook_heading = doc.add_heading("IV. Outlook for Q4", level=1) # Roman numeral, different style
    for run in outlook_heading.runs:
        run.font.name = "Arial Narrow"
        run.font.size = Pt(15)

    p4 = doc.add_paragraph(
        "The outlook for Q4 remains positive, with projected growth of 12-14%. "
        "Focus will be on cost optimization and expanding market reach for Project Titan."
    )
    # Very small font
    for run in p4.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
    p4.paragraph_format.line_spacing = 1.0
    p4.paragraph_format.space_after = Pt(0)

    doc.save(filepath)
    print(f"Document '{filepath}' created successfully with mixed formatting.")

if __name__ == "__main__":
    create_board_report()
