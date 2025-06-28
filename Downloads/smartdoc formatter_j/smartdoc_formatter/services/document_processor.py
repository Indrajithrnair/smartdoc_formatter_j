from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
import os
from typing import Dict, List, Optional, Any, Tuple
import logging
from functools import wraps
import time
from datetime import datetime

from langchain_groq import ChatGroq
from langchain.agents import AgentExecutor, create_openai_functions_agent
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.tools import Tool
from langchain_core.messages import SystemMessage, HumanMessage, AIMessage
from langchain_community.document_loaders import TextLoader
from langchain.memory import ConversationBufferMemory
from langchain_core.runnables import RunnablePassthrough

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.table import _Cell, _Row, _Column, Table

# Configure logging
logger = logging.getLogger(__name__)

def check_environment():
    """Check and validate required environment variables."""
    required_vars = {
        'GROQ_API_KEY': 'Please set your Groq API key in the .env file',
        'FLASK_APP': 'FLASK_APP should be set to app.py',
        'SECRET_KEY': 'Please set a secure SECRET_KEY in the .env file'
    }
    
    missing_vars = []
    for var, message in required_vars.items():
        if not os.getenv(var):
            missing_vars.append(f"{var}: {message}")
    
    if missing_vars:
        error_message = "Missing required environment variables:\n" + "\n".join(missing_vars)
        logger.error(error_message)
        raise ValueError(error_message)

def retry_on_rate_limit(max_retries=3, delay=2):
    """Decorator to handle Groq API rate limits with exponential backoff."""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            retries = 0
            while retries < max_retries:
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    if "rate limit" in str(e).lower():
                        wait_time = delay * (2 ** retries)
                        logger.warning(f"Rate limit hit, retrying in {wait_time} seconds...")
                        time.sleep(wait_time)
                        retries += 1
                    else:
                        raise
            raise Exception("Max retries exceeded for rate limit")
        return wrapper
    return decorator

class DocumentTools:
    """Enhanced tools for document manipulation using python-docx."""
    
    @staticmethod
    def apply_heading(paragraph, level: int):
        """Apply heading style to paragraph."""
        if not 1 <= level <= 6:
            raise ValueError("Heading level must be between 1 and 6")
        paragraph.style = f'Heading {level}'
        
    @staticmethod
    def apply_font_style(run, **kwargs):
        """Apply font styling to a run."""
        if 'bold' in kwargs:
            run.bold = kwargs['bold']
        if 'italic' in kwargs:
            run.italic = kwargs['italic']
        if 'underline' in kwargs:
            run.underline = kwargs['underline']
        if 'strike' in kwargs:
            run.strike = kwargs['strike']
        if 'size' in kwargs:
            run.font.size = Pt(kwargs['size'])
        if 'color' in kwargs:
            rgb = kwargs['color']
            run.font.color.rgb = RGBColor(*rgb)
        if 'name' in kwargs:
            run.font.name = kwargs['name']
            
    @staticmethod
    def set_paragraph_format(paragraph, **kwargs):
        """Set paragraph formatting options."""
        if 'alignment' in kwargs:
            align_map = {
                'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
            paragraph.alignment = align_map.get(kwargs['alignment'].lower())
            
        if 'line_spacing' in kwargs:
            paragraph.paragraph_format.line_spacing = kwargs['line_spacing']
            
        if 'space_before' in kwargs:
            paragraph.paragraph_format.space_before = Pt(kwargs['space_before'])
            
        if 'space_after' in kwargs:
            paragraph.paragraph_format.space_after = Pt(kwargs['space_after'])
            
        if 'first_line_indent' in kwargs:
            paragraph.paragraph_format.first_line_indent = Inches(kwargs['first_line_indent'])
            
        if 'left_indent' in kwargs:
            paragraph.paragraph_format.left_indent = Inches(kwargs['left_indent'])
            
        if 'right_indent' in kwargs:
            paragraph.paragraph_format.right_indent = Inches(kwargs['right_indent'])

    @staticmethod
    def create_table(document, rows: int, cols: int, **kwargs) -> Table:
        """Create and format a table."""
        table = document.add_table(rows=rows, cols=cols)
        
        # Apply table style
        if 'style' in kwargs:
            table.style = kwargs.get('style', 'Table Grid')
            
        # Set table alignment
        if 'alignment' in kwargs:
            align_map = {
                'left': WD_TABLE_ALIGNMENT.LEFT,
                'center': WD_TABLE_ALIGNMENT.CENTER,
                'right': WD_TABLE_ALIGNMENT.RIGHT
            }
            table.alignment = align_map.get(kwargs['alignment'].lower())
            
        # Set column widths
        if 'col_widths' in kwargs:
            for i, width in enumerate(kwargs['col_widths']):
                for cell in table.columns[i].cells:
                    cell.width = Inches(width)
                    
        return table

    @staticmethod
    def format_table_cell(cell: _Cell, **kwargs):
        """Format a table cell."""
        if 'vertical_alignment' in kwargs:
            align_map = {
                'top': WD_CELL_VERTICAL_ALIGNMENT.TOP,
                'center': WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                'bottom': WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            }
            cell.vertical_alignment = align_map.get(kwargs['vertical_alignment'].lower())
            
        if 'background_color' in kwargs:
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{kwargs["background_color"]}"/>'))
            
        if 'margins' in kwargs:
            margins = kwargs['margins']
            cell.margin_top = Cm(margins.get('top', 0))
            cell.margin_bottom = Cm(margins.get('bottom', 0))
            cell.margin_left = Cm(margins.get('left', 0))
            cell.margin_right = Cm(margins.get('right', 0))

    @staticmethod
    def merge_table_cells(table: Table, start_row: int, start_col: int, end_row: int, end_col: int):
        """Merge a range of table cells."""
        if not (0 <= start_row <= end_row < len(table.rows) and 
                0 <= start_col <= end_col < len(table.columns)):
            raise ValueError("Invalid cell range for merging")
            
        # Get the first cell in the range
        merged_cell = table.cell(start_row, start_col)
        
        # Merge the cells
        merged_cell.merge(table.cell(end_row, end_col))
        
        return merged_cell

class DocumentAgent:
    """Enhanced LangChain agent for document processing using Groq."""
    
    def __init__(self):
        """Initialize the document processor with Groq-powered agent."""
        # Check environment variables before initialization
        check_environment()
        
        self.api_key = os.getenv("GROQ_API_KEY")
        
        # Initialize primary and secondary LLMs
        self.primary_llm = self._init_llm("llama3-70b-8192", 0.1, 4096)
        self.secondary_llm = self._init_llm("llama3-8b-8192", 0.1, 2048)
        
        # Initialize tools
        self.tools = self._init_tools()
        
        # Initialize memory with increased context
        self.memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True,
            output_key="output",
            input_key="input"
        )
        
        # Initialize agent
        self.agent = self._init_agent()
        
    def _init_llm(self, model_name: str, temperature: float, max_tokens: int) -> ChatGroq:
        """Initialize a Groq LLM with specified parameters."""
        try:
            from langchain_groq import ChatGroq
            return ChatGroq(
                model_name=model_name,
                temperature=temperature,
                groq_api_key=self.api_key
            )
        except Exception as e:
            logger.error(f"Error initializing Groq LLM: {str(e)}")
            raise

    def _init_tools(self) -> List[Tool]:
        """Initialize enhanced document manipulation tools."""
        return [
            Tool(
                name="apply_heading",
                func=DocumentTools.apply_heading,
                description="Apply heading style (1-6) to a paragraph"
            ),
            Tool(
                name="apply_font_style",
                func=DocumentTools.apply_font_style,
                description="Apply font styling (bold, italic, underline, strike, size, color, name)"
            ),
            Tool(
                name="set_paragraph_format",
                func=DocumentTools.set_paragraph_format,
                description="Set paragraph formatting (alignment, spacing, indentation)"
            ),
            Tool(
                name="create_table",
                func=DocumentTools.create_table,
                description="Create and format a table with specified dimensions and style"
            ),
            Tool(
                name="format_table_cell",
                func=DocumentTools.format_table_cell,
                description="Format table cell (alignment, background, margins)"
            ),
            Tool(
                name="merge_table_cells",
                func=DocumentTools.merge_table_cells,
                description="Merge a range of table cells"
            )
        ]

    def _init_agent(self) -> AgentExecutor:
        """Initialize the LangChain agent with enhanced capabilities."""
        system_message = SystemMessage(content="""
        You are an expert document formatting assistant with deep knowledge of Microsoft Word formatting.
        Your capabilities include:
        1. Comprehensive document structure analysis
        2. Advanced formatting operations:
           - Text styling (font, size, color, emphasis)
           - Paragraph formatting (alignment, spacing, indentation)
           - Table creation and formatting
           - Style application and management
        3. Natural language instruction processing
        4. Document integrity maintenance
        
        Guidelines:
        - Always validate formatting requests for consistency
        - Maintain document structure and hierarchy
        - Provide clear explanations of changes made
        - Handle errors gracefully and provide helpful feedback
        """)
        
        prompt = ChatPromptTemplate.from_messages([
            system_message,
            MessagesPlaceholder(variable_name="chat_history"),
            HumanMessage(content="{input}"),
            MessagesPlaceholder(variable_name="agent_scratchpad"),
        ])

        agent = create_openai_functions_agent(
            llm=self.primary_llm,
            tools=self.tools,
            prompt=prompt
        )
        
        return AgentExecutor(
            agent=agent,
            tools=self.tools,
            memory=self.memory,
            verbose=True,
            handle_parsing_errors=True,
            max_iterations=5
        )

    @retry_on_rate_limit()
    def analyze_document(self, doc: Document) -> Dict[str, Any]:
        """Perform comprehensive document analysis."""
        try:
            # Extract document structure and content
            structure = {
                'paragraphs': len(doc.paragraphs),
                'sections': len(doc.sections),
                'tables': len(doc.tables),
                'styles': list(doc.styles),
                'content': []
            }
            
            # Analyze each paragraph
            for para in doc.paragraphs:
                structure['content'].append({
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal',
                    'alignment': para.alignment,
                    'level': para._p.pPr.numPr.ilvl.val if para._p.pPr and para._p.pPr.numPr else None
                })
            
            # Use secondary LLM for quick analysis
            analysis_chain = LLMChain(
                llm=self.secondary_llm,
                prompt=ChatPromptTemplate.from_messages([
                    SystemMessage(content="Analyze the document structure and provide detailed insights about formatting, hierarchy, and content organization."),
                    HumanMessage(content=str(structure))
                ])
            )
            
            analysis = analysis_chain.run(structure)
            return {
                "status": "success",
                "analysis": analysis,
                "structure": structure,
                "timestamp": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error analyzing document: {str(e)}")
            return {
                "status": "error",
                "message": str(e),
                "timestamp": datetime.utcnow().isoformat()
            }

    @retry_on_rate_limit()
    def process_document(self, doc_path: str, instructions: Optional[str] = None) -> Dict[str, Any]:
        """Process document with enhanced formatting capabilities."""
        try:
            # Load document
            doc = Document(doc_path)
            
            # Analyze document first
            analysis = self.analyze_document(doc)
            if analysis["status"] == "error":
                return analysis
            
            if not instructions:
                return {
                    "status": "success",
                    "message": "Document analyzed successfully",
                    "analysis": analysis
                }
            
            # Process instructions with primary LLM
            result = self.agent.invoke({
                "input": f"""
                Document Analysis: {analysis['analysis']}
                
                Document Structure:
                - Paragraphs: {analysis['structure']['paragraphs']}
                - Tables: {analysis['structure']['tables']}
                - Sections: {analysis['structure']['sections']}
                
                Instructions: {instructions}
                
                Please process these instructions and apply appropriate formatting.
                Provide a detailed summary of changes made.
                """
            })
            
            # Save processed document
            output_path = doc_path.replace('.docx', '_processed.docx')
            doc.save(output_path)
            
            return {
                "status": "success",
                "message": "Document processed successfully",
                "agent_response": result['output'],
                "analysis": analysis,
                "output_path": output_path,
                "timestamp": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            return {
                "status": "error",
                "message": str(e),
                "timestamp": datetime.utcnow().isoformat()
            }

    def reset_memory(self):
        """Reset the agent's conversation memory."""
        self.memory.clear()

class DocumentProcessor:
    def __init__(self):
        """Initialize the document processor with Groq-powered agent."""
        try:
            self.agent = DocumentAgent()
        except Exception as e:
            logger.error(f"Failed to initialize DocumentAgent: {str(e)}")
            raise

    def process_document(self, file_path: str, instructions: Optional[str] = None) -> Dict[str, Any]:
        """
        Process a document with optional formatting instructions.
        
        Args:
            file_path: Path to the document file
            instructions: Optional natural language instructions for formatting
            
        Returns:
            Dict containing processing status and results
        """
        try:
            # Process document using the agent
            result = self.agent.process_document(file_path, instructions)
            
            # Reset agent memory after processing
            self.agent.reset_memory()
            
            return result
            
        except Exception as e:
            logger.error(f"Error in document processing: {str(e)}")
            return {
                "status": "error",
                "message": f"Failed to process document: {str(e)}",
                "timestamp": datetime.utcnow().isoformat()
            } 