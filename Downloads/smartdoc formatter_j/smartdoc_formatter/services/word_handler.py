from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
import mammoth
import os
from typing import Dict, List, Any, Optional, Tuple
import json
import logging

logger = logging.getLogger(__name__)

class DocumentHandler:
    def __init__(self):
        self.document = None
        self.current_file_path = None

    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load a .docx file."""
        try:
            self.document = Document(file_path)
            self.current_file_path = file_path
            return {
                'status': 'success',
                'message': 'Document loaded successfully'
            }
        except Exception as e:
            logger.error(f"Error loading document: {str(e)}")
            return {
                'status': 'error',
                'message': f'Error loading document: {str(e)}'
            }

    def save_document(self, file_path: Optional[str] = None) -> Dict[str, Any]:
        """Save modified document."""
        try:
            save_path = file_path or self.current_file_path
            if not save_path:
                raise ValueError("No file path specified")
            
            self.document.save(save_path)
            return {
                'status': 'success',
                'message': 'Document saved successfully'
            }
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            return {
                'status': 'error',
                'message': f'Error saving document: {str(e)}'
            }

    def extract_text_content(self) -> Dict[str, Any]:
        """Get document text for AI analysis."""
        try:
            content = {
                'paragraphs': [],
                'tables': [],
                'sections': []
            }
            
            # Extract paragraphs
            for para in self.document.paragraphs:
                content['paragraphs'].append({
                    'text': para.text,
                    'style': para.style.name,
                    'runs': [{
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline
                    } for run in para.runs]
                })
            
            # Extract tables
            for table in self.document.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                content['tables'].append(table_data)
            
            # Extract sections
            for section in self.document.sections:
                content['sections'].append({
                    'start_type': section.start_type,
                    'orientation': section.orientation,
                    'page_height': section.page_height,
                    'page_width': section.page_width
                })
            
            return {
                'status': 'success',
                'content': content
            }
        except Exception as e:
            logger.error(f"Error extracting text content: {str(e)}")
            return {
                'status': 'error',
                'message': f'Error extracting text content: {str(e)}'
            }

    def apply_formatting_changes(self, changes_dict: Dict[str, Any]) -> Dict[str, Any]:
        """Apply AI-suggested changes."""
        try:
            # Apply text formatting
            if 'text_formatting' in changes_dict:
                for change in changes_dict['text_formatting']:
                    para_index = change.get('paragraph_index')
                    if para_index is not None and para_index < len(self.document.paragraphs):
                        self.apply_text_formatting(
                            self.document.paragraphs[para_index],
                            change.get('formatting_rules', {})
                        )
            
            # Apply paragraph structure changes
            if 'paragraph_structure' in changes_dict:
                for change in changes_dict['paragraph_structure']:
                    para_index = change.get('paragraph_index')
                    if para_index is not None and para_index < len(self.document.paragraphs):
                        self.modify_paragraph_structure(
                            self.document.paragraphs[para_index],
                            change.get('modifications', {})
                        )
            
            # Apply table formatting
            if 'table_formatting' in changes_dict:
                for change in changes_dict['table_formatting']:
                    table_index = change.get('table_index')
                    if table_index is not None and table_index < len(self.document.tables):
                        self.handle_table_formatting(
                            self.document.tables[table_index],
                            change.get('rules', {})
                        )
            
            # Apply style changes
            if 'style_changes' in changes_dict:
                self.apply_style_changes(self.document, changes_dict['style_changes'])
            
            return {
                'status': 'success',
                'message': 'Formatting changes applied successfully'
            }
        except Exception as e:
            logger.error(f"Error applying formatting changes: {str(e)}")
            return {
                'status': 'error',
                'message': f'Error applying formatting changes: {str(e)}'
            }

    def generate_html_preview(self) -> Dict[str, Any]:
        """Create HTML preview of document using mammoth."""
        try:
            if not self.current_file_path:
                raise ValueError("No document is currently loaded")
            
            with open(self.current_file_path, 'rb') as docx_file:
                result = mammoth.convert_to_html(docx_file)
                messages = result.messages  # Warnings during conversion
                
                return {
                    'status': 'success',
                    'html': result.value,
                    'messages': messages
                }
        except Exception as e:
            logger.error(f"Error generating HTML preview: {str(e)}")
            return {
                'status': 'error',
                'message': f'Error generating HTML preview: {str(e)}'
            }

    def apply_text_formatting(self, paragraph: Any, formatting_rules: Dict[str, Any]) -> None:
        """Apply text formatting to a paragraph."""
        try:
            # Apply font formatting
            if 'font' in formatting_rules:
                font_rules = formatting_rules['font']
                for run in paragraph.runs:
                    if 'name' in font_rules:
                        run.font.name = font_rules['name']
                    if 'size' in font_rules:
                        run.font.size = Pt(font_rules['size'])
                    if 'bold' in font_rules:
                        run.font.bold = font_rules['bold']
                    if 'italic' in font_rules:
                        run.font.italic = font_rules['italic']
                    if 'color' in font_rules:
                        rgb = font_rules['color']
                        run.font.color.rgb = RGBColor(*rgb)
            
            # Apply paragraph formatting
            if 'paragraph' in formatting_rules:
                para_rules = formatting_rules['paragraph']
                if 'alignment' in para_rules:
                    align_map = {
                        'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                        'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                        'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                        'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    }
                    paragraph.alignment = align_map.get(para_rules['alignment'].lower())
                
                if 'line_spacing' in para_rules:
                    paragraph.paragraph_format.line_spacing = para_rules['line_spacing']
                
                if 'space_before' in para_rules:
                    paragraph.paragraph_format.space_before = Pt(para_rules['space_before'])
                
                if 'space_after' in para_rules:
                    paragraph.paragraph_format.space_after = Pt(para_rules['space_after'])
        except Exception as e:
            logger.error(f"Error applying text formatting: {str(e)}")
            raise

    def modify_paragraph_structure(self, paragraph: Any, modifications: Dict[str, Any]) -> None:
        """Modify paragraph structure."""
        try:
            if 'style' in modifications:
                paragraph.style = modifications['style']
            
            if 'level' in modifications:
                paragraph.paragraph_format.outline_level = modifications['level']
            
            if 'indentation' in modifications:
                indent = modifications['indentation']
                if 'left' in indent:
                    paragraph.paragraph_format.left_indent = Inches(indent['left'])
                if 'right' in indent:
                    paragraph.paragraph_format.right_indent = Inches(indent['right'])
                if 'first_line' in indent:
                    paragraph.paragraph_format.first_line_indent = Inches(indent['first_line'])
            
            if 'spacing' in modifications:
                spacing = modifications['spacing']
                if 'before' in spacing:
                    paragraph.paragraph_format.space_before = Pt(spacing['before'])
                if 'after' in spacing:
                    paragraph.paragraph_format.space_after = Pt(spacing['after'])
                if 'line' in spacing:
                    paragraph.paragraph_format.line_spacing = spacing['line']
        except Exception as e:
            logger.error(f"Error modifying paragraph structure: {str(e)}")
            raise

    def handle_table_formatting(self, table: Any, rules: Dict[str, Any]) -> None:
        """Apply formatting to a table."""
        try:
            if 'style' in rules:
                table.style = rules['style']
            
            if 'alignment' in rules:
                align_map = {
                    'left': WD_TABLE_ALIGNMENT.LEFT,
                    'center': WD_TABLE_ALIGNMENT.CENTER,
                    'right': WD_TABLE_ALIGNMENT.RIGHT
                }
                table.alignment = align_map.get(rules['alignment'].lower())
            
            if 'cell_formatting' in rules:
                cell_rules = rules['cell_formatting']
                for row in table.rows:
                    for cell in row.cells:
                        if 'vertical_alignment' in cell_rules:
                            align_map = {
                                'top': WD_CELL_VERTICAL_ALIGNMENT.TOP,
                                'center': WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                                'bottom': WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
                            }
                            cell.vertical_alignment = align_map.get(cell_rules['vertical_alignment'].lower())
                        
                        if 'margins' in cell_rules:
                            margins = cell_rules['margins']
                            if 'top' in margins:
                                cell.margin_top = Inches(margins['top'])
                            if 'bottom' in margins:
                                cell.margin_bottom = Inches(margins['bottom'])
                            if 'left' in margins:
                                cell.margin_left = Inches(margins['left'])
                            if 'right' in margins:
                                cell.margin_right = Inches(margins['right'])
        except Exception as e:
            logger.error(f"Error handling table formatting: {str(e)}")
            raise

    def apply_style_changes(self, document: Any, style_rules: Dict[str, Any]) -> None:
        """Apply style changes to the document."""
        try:
            for style_id, properties in style_rules.items():
                style = document.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
                
                if 'font' in properties:
                    font = properties['font']
                    style.font.name = font.get('name', style.font.name)
                    if 'size' in font:
                        style.font.size = Pt(font['size'])
                    style.font.bold = font.get('bold', style.font.bold)
                    style.font.italic = font.get('italic', style.font.italic)
                
                if 'paragraph' in properties:
                    para = properties['paragraph']
                    if 'alignment' in para:
                        align_map = {
                            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                            'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        }
                        style.paragraph_format.alignment = align_map.get(para['alignment'].lower())
                    
                    if 'line_spacing' in para:
                        style.paragraph_format.line_spacing = para['line_spacing']
        except Exception as e:
            logger.error(f"Error applying style changes: {str(e)}")
            raise

    def extract_document_outline(self) -> Dict[str, Any]:
        """Extract document outline/structure."""
        try:
            outline = []
            current_level = 0
            stack = [(0, outline)]
            
            for paragraph in self.document.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name[-1])
                    item = {
                        'text': paragraph.text,
                        'level': level,
                        'children': []
                    }
                    
                    while level <= stack[-1][0]:
                        stack.pop()
                    
                    stack[-1][1].append(item)
                    stack.append((level, item['children']))
            
            return {
                'status': 'success',
                'outline': outline
            }
        except Exception as e:
            logger.error(f"Error extracting document outline: {str(e)}")
            return {
                'status': 'error',
                'message': f'Error extracting document outline: {str(e)}'
            }

    def identify_content_sections(self) -> Dict[str, Any]:
        """Identify and analyze content sections."""
        try:
            sections = []
            current_section = None
            
            for paragraph in self.document.paragraphs:
                # Start a new section on headings
                if paragraph.style.name.startswith('Heading'):
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        'title': paragraph.text,
                        'level': int(paragraph.style.name[-1]),
                        'content': [],
                        'formatting': self._analyze_paragraph_formatting(paragraph)
                    }
                elif current_section:
                    current_section['content'].append({
                        'text': paragraph.text,
                        'formatting': self._analyze_paragraph_formatting(paragraph)
                    })
            
            # Add the last section
            if current_section:
                sections.append(current_section)
            
            return {
                'status': 'success',
                'sections': sections
            }
        except Exception as e:
            logger.error(f"Error identifying content sections: {str(e)}")
            return {
                'status': 'error',
                'message': f'Error identifying content sections: {str(e)}'
            }

    def analyze_current_formatting(self) -> Dict[str, Any]:
        """Analyze current document formatting."""
        try:
            analysis = {
                'styles': {},
                'paragraphs': [],
                'tables': [],
                'sections': []
            }
            
            # Analyze styles
            for style in self.document.styles:
                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    analysis['styles'][style.name] = {
                        'font': {
                            'name': style.font.name,
                            'size': style.font.size.pt if style.font.size else None,
                            'bold': style.font.bold,
                            'italic': style.font.italic
                        },
                        'paragraph': {
                            'alignment': style.paragraph_format.alignment,
                            'line_spacing': style.paragraph_format.line_spacing
                        }
                    }
            
            # Analyze paragraphs
            for para in self.document.paragraphs:
                analysis['paragraphs'].append(self._analyze_paragraph_formatting(para))
            
            # Analyze tables
            for table in self.document.tables:
                table_analysis = {
                    'style': table.style.name,
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'alignment': table.alignment
                }
                analysis['tables'].append(table_analysis)
            
            # Analyze sections
            for section in self.document.sections:
                section_analysis = {
                    'start_type': section.start_type,
                    'orientation': section.orientation,
                    'page_height': section.page_height,
                    'page_width': section.page_width
                }
                analysis['sections'].append(section_analysis)
            
            return {
                'status': 'success',
                'analysis': analysis
            }
        except Exception as e:
            logger.error(f"Error analyzing current formatting: {str(e)}")
            return {
                'status': 'error',
                'message': f'Error analyzing current formatting: {str(e)}'
            }

    def _analyze_paragraph_formatting(self, paragraph: Any) -> Dict[str, Any]:
        """Helper method to analyze paragraph formatting."""
        try:
            return {
                'style': paragraph.style.name,
                'alignment': paragraph.alignment,
                'line_spacing': paragraph.paragraph_format.line_spacing,
                'space_before': paragraph.paragraph_format.space_before,
                'space_after': paragraph.paragraph_format.space_after,
                'runs': [{
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size.pt if run.font.size else None
                } for run in paragraph.runs]
            }
        except Exception as e:
            logger.error(f"Error analyzing paragraph formatting: {str(e)}")
            return {} 