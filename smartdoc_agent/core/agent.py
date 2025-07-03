from langchain_groq import ChatGroq
from langchain.agents import create_react_agent
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage, AIMessage

from smartdoc_agent.core.tools import (
    analyze_document_structure,
    create_formatting_plan,
    apply_contextual_formatting,
    validate_formatting_result
)
from smartdoc_agent.config import get_groq_api_key

import json
import ast
import traceback
import time
from langchain_core.agents import AgentAction, AgentFinish
from typing import Any # Import Any

class DocumentFormattingAgent:
    def __init__(self, model_name="llama3-70b-8192", temperature=0, progress_callback=None):
        self.progress_callback = progress_callback
        self._emit_progress({"type": "lifecycle", "event": "agent_init_start"})

        self.api_key = get_groq_api_key()

        if not self.api_key or self.api_key in ["YOUR_GROQ_API_KEY_HERE"] or "DUMMY" in self.api_key.upper() :
            self._emit_progress({"type": "warning", "message": f"API key is dummy, placeholder, or missing ('{self.api_key}'). Agent's LLM may fail, or tools requiring LLM will fail."})

        try:
            self.llm = ChatGroq(
                groq_api_key=self.api_key,
                model_name=model_name,
                temperature=temperature
            )
            self._emit_progress({"type": "info", "message": f"Core LLM ({model_name}) for agent reasoning initialized."})
        except Exception as e:
            self._emit_progress({"type": "error", "message": f"Fatal: Failed to initialize core LLM for agent: {e}"})
            raise ValueError(f"Failed to initialize core LLM: {e}")

        self.tools = [
            analyze_document_structure,
            create_formatting_plan,
            apply_contextual_formatting,
            validate_formatting_result
        ]
        self.tool_names = [tool.name for tool in self.tools]

        prompt_template = """
        You are an expert document formatting assistant. Your goal is to help users format their .docx documents.
        You have access to the following tools:

        {tools}

        To use a tool, use the following format:
        Thought: Do I need to use a tool? Yes
        Action: The action to take. Must be one of [{tool_names}]
        Action Input: If the tool takes multiple arguments, provide a VALID JSON string representing a dictionary of these arguments (e.g., "{{\"arg1\": \"value1\", \"arg2\": \"value2\"}}"). Ensure keys and string values are in double quotes. If the tool takes a single simple string argument (like a file path for 'analyze_document_structure'), provide just that string WITHOUT surrounding quotes unless they are part of the path itself.
        Observation: The result of the action.

        When you have a response to say to the Human, or if you do not need to use a tool, you MUST use the format:
        Thought: Do I need to use a tool? No
        Final Answer: [your response here]

        A few important rules:
        1. When a tool provides output (Observation), and you need to use that output as input for a subsequent tool, you MUST use the exact, complete output from the Observation for the relevant parameter, UNLESS a special placeholder is specified below.
        2. Ensure your Action Input is correctly formatted for the tool.
        3. For the 'create_formatting_plan' tool:
           - Its 'document_analysis_json' parameter requires a JSON string.
           - Look at the Observation from the 'analyze_document_structure' tool. This observation is a large JSON string.
           - From this large JSON string, find the 'document_path' string value and the 'summary' object.
           - Create a NEW, smaller JSON string that looks like this: {{"document_path": "<actual_path_here>", "summary": <actual_summary_object_here>}}
           - Use this NEW, smaller JSON string as the value for the 'document_analysis_json' parameter in the Action Input for 'create_formatting_plan'.
           Example:
           Observation: {{"elements": [...], "document_path": "path/doc.docx", "summary": {{"total_elements": 10, "paragraph_count": 5, "heading_count": 2}}, ...other_keys...}}
           Thought: I need to make a plan. I will extract the document_path and summary from the observation and use that for create_formatting_plan.
           Action: create_formatting_plan
           Action Input: {{"document_analysis_json": "{{\"document_path\": \"path/doc.docx\", \"summary\": {{\"total_elements\": 10, \"paragraph_count\": 5, \"heading_count\": 2}}}}", "user_goal": "the user's original goal text"}}
        4. For the 'apply_contextual_formatting' tool:
           - It needs 'doc_path', 'formatting_plan_json', 'document_analysis_json', and 'output_doc_path'.
           - 'doc_path' is the original input document path.
           - For 'formatting_plan_json', you MUST use the literal string placeholder: "$CURRENT_FORMATTING_PLAN".
           - For 'document_analysis_json', you MUST use the literal string placeholder: "$FULL_ORIGINAL_ANALYSIS".
           - 'output_doc_path' is where the formatted document should be saved.
           Example Action Input: {{"doc_path": "doc.docx", "formatting_plan_json": "$CURRENT_FORMATTING_PLAN", "document_analysis_json": "$FULL_ORIGINAL_ANALYSIS", "output_doc_path": "doc_formatted.docx"}}
        5. For the 'validate_formatting_result' tool:
           - It requires these arguments: 'original_doc_analysis_json', 'modified_doc_analysis_json', 'formatting_plan_json', and 'user_goal'.
           - For 'original_doc_analysis_json', you MUST use the literal string placeholder: "$FULL_ORIGINAL_ANALYSIS".
           - For 'modified_doc_analysis_json', you MUST use the literal string placeholder: "$FULL_MODIFIED_ANALYSIS" (this refers to the analysis of the document *after* formatting has been applied by `apply_contextual_formatting` and the modified document re-analyzed by `analyze_document_structure`).
           - For 'formatting_plan_json', you MUST use the literal string placeholder: "$CURRENT_FORMATTING_PLAN".
           - 'user_goal' is the original user goal text provided at the start of the task.
           Example Action Input for 'validate_formatting_result':
           {{"original_doc_analysis_json": "$FULL_ORIGINAL_ANALYSIS", "modified_doc_analysis_json": "$FULL_MODIFIED_ANALYSIS", "formatting_plan_json": "$CURRENT_FORMATTING_PLAN", "user_goal": "The user's original request text here"}}

        Begin!

        Previous conversation history:
        {chat_history}

        New input: {input}
        {agent_scratchpad}
        """

        self.prompt = ChatPromptTemplate.from_template(prompt_template).partial(
            tools="\n".join([f"{tool.name}: {tool.description}" for tool in self.tools]),
            tool_names=", ".join([tool.name for tool in self.tools]),
        )

        self.agent = create_react_agent(
            llm=self.llm,
            tools=self.tools,
            prompt=self.prompt
        )

        self.chat_history = []
        self.full_original_analysis_json = None
        self.current_formatting_plan_json = None
        self.full_modified_analysis_json = None
        self.current_doc_path = None
        self.current_output_doc_path = None
        self.PLACEHOLDER_ORIGINAL_ANALYSIS = "$FULL_ORIGINAL_ANALYSIS"
        self.PLACEHOLDER_MODIFIED_ANALYSIS = "$FULL_MODIFIED_ANALYSIS"
        self.PLACEHOLDER_FORMATTING_PLAN = "$CURRENT_FORMATTING_PLAN"

    def _emit_progress(self, data: dict):
        """Helper to call the progress callback if it exists."""
        # print(f"AGENT PROGRESS: {data}") # Optional: for local debugging
        if self.progress_callback:
            try:
                self.progress_callback(data)
            except Exception as e:
                print(f"Error in progress_callback: {e}")

    def _substitute_placeholders(self, action_input: dict) -> dict:
        processed_input = {}
        for key, value in action_input.items():
            if value == self.PLACEHOLDER_ORIGINAL_ANALYSIS:
                if self.full_original_analysis_json:
                    processed_input[key] = self.full_original_analysis_json
                    self._emit_progress({"type": "debug", "message": f"Substituted {self.PLACEHOLDER_ORIGINAL_ANALYSIS} for key '{key}'"})
                else:
                    raise ValueError(f"Agent tried to use placeholder {self.PLACEHOLDER_ORIGINAL_ANALYSIS} but no original analysis is stored.")
            elif value == self.PLACEHOLDER_MODIFIED_ANALYSIS:
                if self.full_modified_analysis_json:
                    processed_input[key] = self.full_modified_analysis_json
                    self._emit_progress({"type": "debug", "message": f"Substituted {self.PLACEHOLDER_MODIFIED_ANALYSIS} for key '{key}'"})
                else:
                    raise ValueError(f"Agent tried to use placeholder {self.PLACEHOLDER_MODIFIED_ANALYSIS} but no modified analysis is stored.")
            elif value == self.PLACEHOLDER_FORMATTING_PLAN:
                if self.current_formatting_plan_json:
                    processed_input[key] = self.current_formatting_plan_json
                    self._emit_progress({"type": "debug", "message": f"Substituted {self.PLACEHOLDER_FORMATTING_PLAN} for key '{key}'"})
                else:
                    raise ValueError(f"Agent tried to use placeholder {self.PLACEHOLDER_FORMATTING_PLAN} but no plan is stored.")
            else:
                processed_input[key] = value
        return processed_input

    def _execute_tool(self, tool_name: str, tool_input_from_llm: Any) -> str:
        self._emit_progress({"type": "tool_start", "name": tool_name, "input_preview": str(tool_input_from_llm)[:100] + "..."})
        payload_value_for_tool: Any = None

        if isinstance(tool_input_from_llm, str):
            parsed_llm_input_dict = None
            try:
                parsed_llm_input_dict = json.loads(tool_input_from_llm)
                self._emit_progress({"type": "debug", "message": f"_execute_tool: Parsed LLM string input as JSON to dict."})
                payload_value_for_tool = self._substitute_placeholders(parsed_llm_input_dict)
            except json.JSONDecodeError:
                payload_value_for_tool = tool_input_from_llm.strip('"').strip("'")
                self._emit_progress({"type": "debug", "message": f"_execute_tool: LLM string input treated as simple string (JSON parse failed, after strip): {str(payload_value_for_tool)[:100]}..."})

        elif isinstance(tool_input_from_llm, dict):
            payload_value_for_tool = self._substitute_placeholders(tool_input_from_llm)
            self._emit_progress({"type": "debug", "message": f"_execute_tool: LLM dict input. After subs: {str(payload_value_for_tool)[:100]}..."})
        else:
            error_msg = f"LLM provided invalid type for Action Input: {type(tool_input_from_llm)}. Value: {str(tool_input_from_llm)[:200]}"
            self._emit_progress({"type": "error", "message": f"_execute_tool: {error_msg}"})
            return json.dumps({"error": error_msg})

        args_for_langchain_invoke = {"tool_input": payload_value_for_tool}

        selected_tool = next((t for t in self.tools if t.name == tool_name), None)
        if not selected_tool:
            error_msg = f"Error: Tool '{tool_name}' not found. Available tools: {[t.name for t in self.tools]}"
            self._emit_progress({"type": "error", "message": error_msg})
            return error_msg

        observation = ""
        try:
            self._emit_progress({"type": "debug", "message": f"Invoking tool '{selected_tool.name}' with args for invoke: {str(args_for_langchain_invoke)[:200]}..."})
            observation = selected_tool.invoke(args_for_langchain_invoke)
            self._emit_progress({"type": "tool_end", "name": tool_name, "observation_preview": str(observation)[:100] + "..."})
        except Exception as e:
            error_msg_tool_exec = f"Error executing tool {tool_name}: {str(e)}"
            self._emit_progress({"type": "error", "message": error_msg_tool_exec, "details": traceback.format_exc()})
            traceback.print_exc() # Also print to server log for immediate visibility
            return error_msg_tool_exec # Return error message as observation

        try:
            # Attempt to parse observation to see if it's JSON (many tools return JSON strings)
            # This is for storing specific data structures, not for general logging.
            parsed_observation_for_storage = json.loads(observation) if isinstance(observation, str) else observation

            if tool_name == "analyze_document_structure" and isinstance(parsed_observation_for_storage, dict) and "document_path" in parsed_observation_for_storage:
                analyzed_path = parsed_observation_for_storage["document_path"]
                if analyzed_path == self.current_doc_path:
                    self.full_original_analysis_json = observation
                    self._emit_progress({"type": "data_store", "variable": "full_original_analysis_json", "length": len(observation) if observation else 0})
                elif analyzed_path == self.current_output_doc_path:
                    self.full_modified_analysis_json = observation
                    self._emit_progress({"type": "data_store", "variable": "full_modified_analysis_json", "length": len(observation) if observation else 0})
            elif tool_name == "create_formatting_plan":
                # Ensure what's stored is the JSON string if the tool returned a string
                if isinstance(observation, str):
                    self.current_formatting_plan_json = observation
                    self._emit_progress({"type": "data_store", "variable": "current_formatting_plan_json", "length": len(observation) if observation else 0})
                elif isinstance(parsed_observation_for_storage, (dict, list)): # If tool already parsed it (should not based on current tools.py)
                    self.current_formatting_plan_json = json.dumps(parsed_observation_for_storage) # Store as string
                    self._emit_progress({"type": "data_store", "variable": "current_formatting_plan_json", "length": len(self.current_formatting_plan_json)})
        except json.JSONDecodeError:
            self._emit_progress({"type":"debug", "message":f"Observation from {tool_name} was not valid JSON, not performing special storage: {observation[:100]}"})
        except Exception as e: # Catch other errors during storage logic
            self._emit_progress({"type":"error", "message":f"Error processing/storing observation from {tool_name}: {e}"})

        return str(observation)

    def run(self, user_query: str, document_path: str, output_document_path: str = None):
        self._emit_progress({"type": "lifecycle", "event": "agent_run_start", "user_query": user_query, "doc_path": document_path})
        self.current_doc_path = document_path
        if not output_document_path:
            self.current_output_doc_path = document_path.replace(".docx", "_formatted.docx")
        else:
            self.current_output_doc_path = output_document_path

        self.full_original_analysis_json = None
        self.current_formatting_plan_json = None
        self.full_modified_analysis_json = None

        agent_initial_input = (
            f"User query: '{user_query}'.\n"
            f"The document to work on is: '{self.current_doc_path}'.\n"
            f"The output path for formatted document is: '{self.current_output_doc_path}'.\n"
            "Follow the general workflow: analyze original, create plan, apply plan, analyze modified, validate. Use placeholders like $FULL_ORIGINAL_ANALYSIS where instructed."
        )

        if not self.chat_history: # Initialize with system message if empty
            self.chat_history.append(HumanMessage(content=agent_initial_input))
        else: # Append if continuing a conversation (not fully supported by this run structure yet)
            self.chat_history.append(HumanMessage(content=agent_initial_input))

        intermediate_steps = []
        max_iterations = 12

        for i in range(max_iterations):
            self._emit_progress({"type": "iteration_start", "iteration": i + 1, "max_iterations": max_iterations})
            try:
                current_input_for_llm = agent_initial_input if i == 0 else "What is the next logical step based on the previous actions and observations?"

                current_agent_input_dict = {
                    "input": current_input_for_llm,
                    "intermediate_steps": intermediate_steps,
                    "chat_history": self.chat_history
                }

                self._emit_progress({"type": "llm_call_start", "purpose": "action_planning", "input_preview": current_input_for_llm[:100]+"..."})
                llm_call_start_time = time.time()
                llm_output = self.agent.invoke(current_agent_input_dict)
                llm_call_duration = time.time() - llm_call_start_time
                self._emit_progress({"type": "llm_call_end", "duration_seconds": round(llm_call_duration,2) , "output_type": str(type(llm_output))})

                if isinstance(llm_output, AgentFinish):
                    final_answer = llm_output.return_values.get("output", "No final answer from AgentFinish.")
                    self._emit_progress({"type": "agent_finish", "final_answer": final_answer})
                    self.chat_history.append(AIMessage(content=final_answer))
                    return final_answer

                if not isinstance(llm_output, AgentAction):
                    final_answer = str(llm_output)
                    if hasattr(llm_output, 'log'):
                        self._emit_progress({"type": "debug", "message": f"LLM Log: {llm_output.log}"})

                    if "Final Answer:" in final_answer :
                         final_answer = final_answer.split("Final Answer:")[-1].strip()
                         self._emit_progress({"type": "agent_finish", "final_answer": final_answer, "reason": "LLM direct final answer"})
                         self.chat_history.append(AIMessage(content=final_answer))
                         return final_answer
                    else:
                        error_msg = f"LLM output was not AgentAction or AgentFinish. Type: {type(llm_output)}, Output: {str(llm_output)[:500]}"
                        self._emit_progress({"type": "error", "message": error_msg})
                        self.chat_history.append(AIMessage(content=f"Agent produced unexpected output type: {type(llm_output)}"))
                        return f"Agent error: Unexpected output type from LLM. Last output: {str(llm_output)[:500]}"

                agent_action = llm_output
                tool_name = agent_action.tool
                tool_input_from_llm = agent_action.tool_input # This is what LLM provides

                self._emit_progress({"type": "debug", "message": f"LLM Action: {tool_name}, LLM Action Input (raw): {str(tool_input_from_llm)[:200]}..."})

                observation_str = self._execute_tool(tool_name, tool_input_from_llm)

                # Truncate very long observations for the scratchpad to manage context window
                observation_str_for_scratchpad = observation_str
                MAX_OBS_LEN_SCRATCHPAD = 1000
                if len(observation_str_for_scratchpad) > MAX_OBS_LEN_SCRATCHPAD:
                    truncated_len = len(observation_str_for_scratchpad)
                    observation_str_for_scratchpad = observation_str_for_scratchpad[:MAX_OBS_LEN_SCRATCHPAD] + f"... (truncated, original length {truncated_len})"

                intermediate_steps.append((agent_action, observation_str_for_scratchpad))

            except Exception as e:
                error_message = f"Error in agent loop iteration {i+1}: {e}"
                self._emit_progress({"type": "error", "message": error_msg, "details": traceback.format_exc()})
                traceback.print_exc()
                self.chat_history.append(AIMessage(content=f"Encountered an error: {error_message}"))
                return f"An error occurred in the agent loop: {e}"

        final_message = "Agent reached maximum iterations."
        self._emit_progress({"type": "lifecycle", "event": "agent_run_max_iterations", "message": final_message})
        self.chat_history.append(AIMessage(content=final_message))
        return final_message


if __name__ == '__main__':
    from docx import Document as PythonDocXDocument
    sample_doc_path = "sample_doc_for_agent.docx"
    sample_output_path = "sample_doc_for_agent_formatted.docx"

    def cli_progress_callback(data: dict):
        print(f"CLI_PROGRESS: {json.dumps(data)}")

    doc = PythonDocXDocument()
    doc.add_heading('Agent Test Document', level=0)
    doc.add_paragraph('This is paragraph one. It needs some formatting.')
    doc.add_heading('Section Alpha', level=1)
    doc.add_paragraph('This is paragraph two, under Section Alpha. It is a bit messy.')
    doc.save(sample_doc_path)
    print(f"Created '{sample_doc_path}' for agent testing.")

    try:
        formatter_agent = DocumentFormattingAgent(progress_callback=cli_progress_callback)

        user_request = "Please make this document look more professional and fix inconsistencies."
        print(f"\n--- Running agent with query: '{user_request}' on '{sample_doc_path}' ---")

        result = formatter_agent.run(
            user_query=user_request,
            document_path=sample_doc_path,
            output_document_path=sample_output_path
        )

        print("\n--- Agent Result ---")
        print(result)

    except ValueError as ve:
        print(f"Setup Error: {ve}")
        if "GROQ_API_KEY is not set" in str(ve):
             print("Please ensure your GROQ_API_KEY is set in .env at the project root.")
    except Exception as ex:
        print(f"An unexpected error occurred: {ex}")
        traceback.print_exc()
    finally:
        import os
        if os.path.exists(sample_doc_path):
            os.remove(sample_doc_path)
        if os.path.exists(sample_output_path):
            os.remove(sample_output_path)
        print(f"\nCleaned up '{sample_doc_path}' and '{sample_output_path}'.")
