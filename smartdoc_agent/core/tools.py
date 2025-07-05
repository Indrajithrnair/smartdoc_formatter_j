import json
import ast # For safely evaluating string literals
from langchain.tools import tool
from smartdoc_agent.utils.document_utils import load_document, save_document, get_document_analysis
# We will properly implement these tools later. For now, they are stubs.

@tool
def analyze_document_structure(tool_input: str | dict) -> str:
    """
    Analyzes the structure and formatting of a .docx document.
    Input can be a string path, or a dictionary like {'doc_path': 'path/to/doc.docx'},
    or a string representation of such a dictionary e.g. \"{'doc_path': 'path/to/doc.docx'}\".
    Returns a JSON string containing detailed analysis of document elements,
    including paragraphs, headings, styles, fonts, and alignments.
    """
    doc_path_str = None
    # The 'tool_input' is now the direct payload prepared by DocumentFormattingAgent._execute_tool
    # It's either a simple string (e.g. path) or a dictionary of arguments.


    doc_path_str = None
    # If input is a stringified dict, parse it
    if isinstance(tool_input, str):
        try:
            # Try to parse as dict
            parsed = ast.literal_eval(tool_input)
            if isinstance(parsed, dict) and "doc_path" in parsed:
                doc_path_str = parsed["doc_path"]
            else:
                doc_path_str = tool_input
        except Exception:
            doc_path_str = tool_input
    elif isinstance(tool_input, dict):
        doc_path_str = tool_input.get("doc_path")

    if not doc_path_str:
        error_msg = f"analyze_document_structure: 'doc_path' not found in tool_input: {tool_input}"
        print(f"Tool Error: {error_msg}")
        return json.dumps({"error": error_msg})

    print(f"Tool: analyze_document_structure called. Effective doc_path='{doc_path_str}' (Original agent tool_input was: {tool_input})")

    try:
        doc = load_document(doc_path_str)
        analysis_data = get_document_analysis(doc)

        analysis_data["document_path"] = doc_path_str # Store the actual path used
        analysis_data["summary"] = {
            "total_elements": len(analysis_data.get("elements", [])),
            "paragraph_count": sum(1 for el in analysis_data.get("elements", []) if el.get("type") == "paragraph"),
            "heading_count": sum(1 for el in analysis_data.get("elements", []) if el.get("type") == "heading"),
        }

        return json.dumps(analysis_data, indent=2)
    except Exception as e:
        error_msg = f"Error in analyze_document_structure with doc_path '{doc_path_str}': {str(e)}"
        print(f"Tool: analyze_document_structure error: {error_msg}")
        return json.dumps({"error": error_msg, "input_doc_path": doc_path_str})

from smartdoc_agent.config import get_groq_api_key
from langchain_groq import ChatGroq

# Initialize the LLM for this tool if needed, or pass it in / access from an agent context
# For simplicity here, we might initialize it directly or assume it's available globally/via config.

groq_api_key_for_tools = get_groq_api_key()
llm = None # Initialize llm to None

if groq_api_key_for_tools and groq_api_key_for_tools not in [
    "DUMMY_KEY_PROJECT_ROOT",
    "DUMMY_KEY_FOR_TESTING_CLI_FLOW",
    "YOUR_GROQ_API_KEY_HERE", # Standard placeholder
    "DUMMY_KEY_DO_NOT_USE_FOR_REAL_CALLS" # My specific dummy key
]:
    try:
        print(f"[core/tools.py] Attempting to initialize ChatGroq with API key: '{groq_api_key_for_tools[:10]}...'") # Print part of the key for confirmation
        llm = ChatGroq(groq_api_key=groq_api_key_for_tools, model_name="llama3-70b-8192", temperature=0) # Updated model name
        print(f"[core/tools.py] ChatGroq LLM initialized successfully for tools.")
    except Exception as e:
        llm = None # Ensure llm is None if initialization fails
        print(f"❌ [core/tools.py] Failed to initialize ChatGroq LLM: {e}. LLM-dependent tools may not function fully.")
else:
    llm = None # Ensure llm is None
    print(f"ℹ️ [core/tools.py] LLM not initialized due to missing, placeholder, or dummy API key ('{groq_api_key_for_tools}'). LLM-dependent tools will use fallback or error.")


@tool
def create_formatting_plan(tool_input: str | dict) -> str:
    """
    Creates a strategic, structured formatting plan (JSON) based on document analysis and a user goal.
    Input can be a dict {'document_analysis_json': '...', 'user_goal': '...'} or its string representation.
    The agent's _execute_tool method ensures 'tool_input' is this dictionary.
    Uses an LLM to interpret the analysis and goal, then generate actionable steps.
    The document_analysis_json is a JSON string from the 'analyze_document_structure' tool.
    The user_goal is a string describing the desired outcome (e.g., "make professional").
    Returns a JSON string representing the list of formatting actions.
    """
    if not isinstance(tool_input, dict):
        error_msg = f"create_formatting_plan expects a dictionary, got {type(tool_input)}. Value: {str(tool_input)[:200]}"
        print(f"Tool Error: {error_msg}")
        return json.dumps({"error": error_msg})

    document_analysis_json_val = tool_input.get("document_analysis_json")
    user_goal = tool_input.get("user_goal")

    if document_analysis_json_val is None or user_goal is None:
        error_msg = f"create_formatting_plan missing 'document_analysis_json' or 'user_goal'. Received keys: {list(tool_input.keys())}"
        print(f"Tool Error: {error_msg}")
        return json.dumps({"error": error_msg, "received_input": tool_input})

    print(f"Tool: create_formatting_plan called with user_goal='{user_goal}'. Effective document_analysis_json (type {type(document_analysis_json_val)}): {str(document_analysis_json_val)[:100]}...")

    if not llm:
        return json.dumps({"error": "LLM not initialized for create_formatting_plan."})

    parsed_analysis_data = None
    if isinstance(document_analysis_json_val, str): # Changed to _val
        try:
            parsed_analysis_data = json.loads(document_analysis_json_val) # Changed to _val
        except json.JSONDecodeError:
            return json.dumps({"error": f"Invalid JSON string for document_analysis_json: {document_analysis_json_val}"}) # Changed to _val
    elif isinstance(document_analysis_json_val, dict): # Changed to _val
        parsed_analysis_data = document_analysis_json_val # Changed to _val
    else:
        return json.dumps({"error": f"document_analysis_json must be a JSON string or a dictionary, got {type(document_analysis_json_val)}."}) # Changed to _val

    if not parsed_analysis_data:
        return json.dumps({"error": "Failed to obtain parsed_analysis_data."})

    if "error" in parsed_analysis_data:
        return json.dumps({"error": f"Provided document analysis itself contains an error: {parsed_analysis_data['error']}"})

    # For brevity in the prompt, use the (already summarized) parsed_analysis_data
    # The ReAct prompt already guides the LLM to create a summary for this tool.
    analysis_summary_for_prompt = {
        "document_path": parsed_analysis_data.get("document_path"), # Should exist if from analyze_document_structure
        "summary": parsed_analysis_data.get("summary"),       # Should exist if from analyze_document_structure
         # elements_sample is not strictly needed if summary is good enough for planning
    }
    if not analysis_summary_for_prompt.get("document_path") or not analysis_summary_for_prompt.get("summary"):
        return json.dumps({"error": "Summarized analysis for planning is missing 'document_path' or 'summary'.", "received_analysis": parsed_analysis_data})

    prompt = f"""
You are an expert document formatting assistant. Your task is to create a detailed, step-by-step formatting plan
in JSON format based on a document analysis and a user's desired goal.

User Goal: "{user_goal}"

Document Analysis Summary:
{json.dumps(analysis_summary_for_prompt, indent=2)}

Based on the user's goal and the document analysis, generate a JSON list of specific, actionable formatting commands.
Each command in the list should be a JSON object with an "action" key and other keys relevant to that action.
Possible actions include (but are not limited to):
- "set_font": For changing font of elements. Needs "scope" (e.g., "all_paragraphs", "headings_level_1", "paragraph_index_N"), "font_name", "size" (in pt). Optional: "bold", "italic".
- "set_heading_style": For defining styles for heading levels. Needs "level" (e.g., 1, 2), "font_name", "size", "bold". Optional: "italic", "keep_with_next".
- "set_paragraph_spacing": For paragraph spacing. Needs "scope", "spacing_before" (in pt), "spacing_after" (in pt), "line_spacing" (e.g., 1.0, 1.5).
- "set_alignment": For text alignment. Needs "scope", "alignment" (e.g., "LEFT", "CENTER", "RIGHT", "JUSTIFY").
- "ensure_consistent_style": For a specific style. Needs "style_name", and desired properties like "font_name", "size".
- "fix_font_inconsistencies": Identify and unify fonts. Needs "target_font_name", "target_font_size".
- "apply_theme": If a general theme is requested (e.g., "modern", "academic"). This is more abstract and might translate to multiple specific actions.

Example of a desired JSON output format:
[
  {{"action": "set_font", "scope": "all_body_paragraphs", "font_name": "Calibri", "size": 12}},
  {{"action": "set_heading_style", "level": 1, "font_name": "Calibri Light", "size": 18, "bold": true, "spacing_after": 12}},
  {{"action": "set_paragraph_spacing", "scope": "all_paragraphs", "spacing_after": 6, "line_spacing": 1.15}},
  {{"action": "set_alignment", "scope": "headings_level_1", "alignment": "LEFT"}}
]

IMPORTANT: Only output the JSON list of actions. Do not include any other text, explanations, or markdown formatting.
The output must be a valid JSON array of objects.
If the user goal is vague, make reasonable assumptions for a standard professional document.
If the analysis shows significant inconsistencies, prioritize fixing those.
Focus on common professional formatting standards.
"""

    try:
        response = llm.invoke(prompt)
        generated_plan_str = response.content

        # Validate if the output is valid JSON
        try:
            json.loads(generated_plan_str) # Try parsing
            return generated_plan_str # Return the string if it's valid JSON
        except json.JSONDecodeError as json_err:
            error_message = f"LLM returned invalid JSON for the plan. Error: {json_err}. LLM Output: {generated_plan_str}"
            print(error_message)
            return json.dumps({"error": "LLM generated invalid JSON plan.", "details": error_message})

    except Exception as e:
        error_message = f"Error during LLM call in create_formatting_plan: {e}"
        print(error_message)
        return json.dumps({"error": error_message})

from smartdoc_agent.utils.document_utils import (
    load_document, save_document, get_document_analysis, # get_document_analysis is needed here
    apply_set_font_action, apply_set_heading_style_action,
    apply_set_paragraph_spacing_action, apply_set_alignment_action,
    apply_fix_font_inconsistencies_action # Added new action handler
)

# ... (other tools like analyze_document_structure, create_formatting_plan remain unchanged from previous step)


@tool
def apply_contextual_formatting(tool_input: str | dict) -> str:
    """
    Applies the formatting plan (JSON string) to the document.
    Input can be a dict or string representation of a dict with keys:
    'doc_path', 'formatting_plan_json', 'document_analysis_json', 'output_doc_path'.
    The document_analysis_json is required to correctly map plan scopes.
    Saves the modified document to output_doc_path.
    Returns a status message, or an error message if issues occur.
    """
    if not isinstance(tool_input, dict):
        error_msg = f"apply_contextual_formatting expects a dictionary, got {type(tool_input)}. Value: {str(tool_input)[:200]}"
        print(f"Tool Error: {error_msg}")
        return json.dumps({"error": error_msg})

    doc_path = tool_input.get("doc_path")
    formatting_plan_json_val = tool_input.get("formatting_plan_json")
    document_analysis_json_val = tool_input.get("document_analysis_json")
    output_doc_path = tool_input.get("output_doc_path")

    required_keys = ['doc_path', 'formatting_plan_json', 'document_analysis_json', 'output_doc_path']
    if not all(tool_input.get(key) is not None for key in required_keys):
        missing_keys = [key for key in required_keys if tool_input.get(key) is None]
        error_msg = f"apply_contextual_formatting missing required arguments: {missing_keys}. Received keys: {list(tool_input.keys())}"
        print(f"Tool Error: {error_msg}")
        return json.dumps({"error": error_msg, "received_input": tool_input})

    print(f"Tool: apply_contextual_formatting called for doc_path='{doc_path}', output_doc_path='{output_doc_path}'.")
    print(f"  Effective formatting_plan_json (type {type(formatting_plan_json_val)}): {str(formatting_plan_json_val)[:100]}...")
    print(f"  Effective document_analysis_json (type {type(document_analysis_json_val)}): {str(document_analysis_json_val)[:100]}...")

    try:
        # formatting_plan_json_val should be a JSON string (substituted by agent if placeholder was used)
        if not isinstance(formatting_plan_json_val, str): # Check if it's a string first
            # If it's not a string, it might be an already parsed list/dict from LLM, or an error
            if isinstance(formatting_plan_json_val, list): # LLM might give a list of dicts
                plan_actions = formatting_plan_json_val
            else: # Unexpected type
                return json.dumps({"error": f"formatting_plan_json must be a JSON string or a list, got {type(formatting_plan_json_val)}."})
        else: # It is a string, try to load
            plan_actions = json.loads(formatting_plan_json_val)

        if not isinstance(plan_actions, list): # After potential loading, check if it's a list
            return json.dumps({"error": "Formatting plan (after parsing) is not a valid JSON list."})

        # Check if the plan itself indicates an error
        if plan_actions and isinstance(plan_actions[0], dict) and "error" in plan_actions[0]:
             return json.dumps({"error": f"Formatting plan indicates an error from a previous step: {plan_actions[0]['error']}"})
    except json.JSONDecodeError: # Only if formatting_plan_json_val was a string and failed to parse
        return json.dumps({"error": f"Invalid JSON format for formatting_plan_json string: {formatting_plan_json_val}"})
    except Exception as e: # Catch other potential errors during plan processing
        return json.dumps({"error": f"Error processing formatting plan: {str(e)}"})


    # --- Document Analysis JSON Handling ---
    # document_analysis_json_val should be a JSON string (substituted by agent if placeholder was used)
    if document_analysis_json_val == "$FULL_ORIGINAL_ANALYSIS": # This check is now mostly for debugging the agent's substitution
        # This case should ideally be handled by the agent substituting the placeholder.
        # If the tool receives this, it means substitution didn't happen as expected upstream.
        return json.dumps({"error": "Placeholder $FULL_ORIGINAL_ANALYSIS was passed to tool, but substitution should have occurred in agent."})

    parsed_analysis_data = None
    if isinstance(document_analysis_json_val, str):
        try:
            parsed_analysis_data = json.loads(document_analysis_json_val)
        except json.JSONDecodeError:
            return json.dumps({"error": f"Invalid JSON string for document_analysis_json: {document_analysis_json_val}"})
    elif isinstance(document_analysis_json_val, dict):
        parsed_analysis_data = document_analysis_json_val
    else:
        return json.dumps({"error": f"document_analysis_json must be a JSON string or dict, got {type(document_analysis_json_val)}."})

    if "error" in parsed_analysis_data:
         return json.dumps({"error": f"Provided document analysis itself contains an error: {parsed_analysis_data['error']}"})
    elements_details = parsed_analysis_data.get("elements", [])
    if not elements_details:
        return json.dumps({"error": "Document analysis data is missing 'elements' list. Full analysis is required.", "received_analysis_keys": list(parsed_analysis_data.keys())})

    try:
        doc = load_document(doc_path)

        action_handlers = {
            "set_font": apply_set_font_action,
            "set_heading_style": apply_set_heading_style_action,
            "set_paragraph_spacing": apply_set_paragraph_spacing_action,
            "set_alignment": apply_set_alignment_action,
            "fix_font_inconsistencies": apply_fix_font_inconsistencies_action,
            # Add other handlers here as they are implemented
            # "ensure_consistent_style": apply_ensure_consistent_style_action,
        }

        applied_actions_count = 0
        skipped_actions_count = 0

        for action in plan_actions:
            action_type = action.get("action")
            handler = action_handlers.get(action_type)

            if handler:
                try:
                    # Pass the document object, the detailed element list from analysis, and the action itself
                    handler(doc, elements_details, action)
                    applied_actions_count += 1
                except Exception as e:
                    print(f"Error applying action {action}: {e}")
                    # Optionally, collect these errors to return in the result
                    skipped_actions_count +=1
            else:
                print(f"Warning: No handler found for action type '{action_type}'. Skipping.")
                skipped_actions_count +=1

        save_document(doc, output_doc_path)
        return json.dumps({
            "status": "success",
            "message": f"Formatting applied. Document saved to {output_doc_path}.",
            "actions_applied": applied_actions_count,
            "actions_skipped_or_failed": skipped_actions_count
        })
    except Exception as e:
        return json.dumps({"error": f"Error in apply_contextual_formatting: {str(e)}"})

@tool
def validate_formatting_result(tool_input: str | dict) -> str:
    """
    Validates if formatting goals were achieved.
    Input can be a dict or string representation of a dict with keys:
    'original_doc_analysis_json', 'modified_doc_analysis_json', 'formatting_plan_json', 'user_goal'.
    The agent's _execute_tool method ensures 'tool_input' is this dictionary.
    Compares document analyses against the plan and user goal.
    Uses an LLM for qualitative assessment. Returns a JSON validation report.
    """
    tool_name = "validate_formatting_result"
    if not isinstance(tool_input, dict):
        error_msg = f"{tool_name} expects a dictionary, got {type(tool_input)}. Value: {str(tool_input)[:200]}"
        print(f"Tool Error: {error_msg}")
        return json.dumps({"error": error_msg})

    original_doc_analysis_json_val = tool_input.get("original_doc_analysis_json")
    modified_doc_analysis_json_val = tool_input.get("modified_doc_analysis_json")
    formatting_plan_json_val = tool_input.get("formatting_plan_json")
    user_goal = tool_input.get("user_goal")

    required_keys = ['original_doc_analysis_json', 'modified_doc_analysis_json', 'formatting_plan_json', 'user_goal']
    if not all(tool_input.get(key) is not None for key in required_keys):
        missing_keys = [key for key in required_keys if tool_input.get(key) is None]
        error_msg = f"{tool_name} missing required arguments: {missing_keys}. Received keys: {list(tool_input.keys())}"
        print(f"Tool Error: {error_msg}")
        return json.dumps({"error": error_msg, "received_input": tool_input})

    print(f"Tool: {tool_name} called with user_goal='{user_goal}'.")
    print(f"  Effective original_doc_analysis_json (type {type(original_doc_analysis_json_val)}): {str(original_doc_analysis_json_val)[:100]}...")
    print(f"  Effective modified_doc_analysis_json (type {type(modified_doc_analysis_json_val)}): {str(modified_doc_analysis_json_val)[:100]}...")
    print(f"  Effective formatting_plan_json (type {type(formatting_plan_json_val)}): {str(formatting_plan_json_val)[:100]}...")

    if not llm:
        return json.dumps({"error": f"LLM not initialized for {tool_name}."})

    try:
        # original_doc_analysis_json_val handling
        if isinstance(original_doc_analysis_json_val, str):
            original_analysis = json.loads(original_doc_analysis_json_val)
        elif isinstance(original_doc_analysis_json_val, dict):
            original_analysis = original_doc_analysis_json_val
        else:
            return json.dumps({"error": f"original_doc_analysis_json must be JSON string or dict, got {type(original_doc_analysis_json_val)}."})

        # modified_doc_analysis_json_val handling
        if isinstance(modified_doc_analysis_json_val, str):
            modified_analysis = json.loads(modified_doc_analysis_json_val)
        elif isinstance(modified_doc_analysis_json_val, dict):
            modified_analysis = modified_doc_analysis_json_val
        else:
            return json.dumps({"error": f"modified_doc_analysis_json must be JSON string or dict, got {type(modified_doc_analysis_json_val)}."})

        # formatting_plan_json_val handling
        if isinstance(formatting_plan_json_val, str):
            plan = json.loads(formatting_plan_json_val)
        elif isinstance(formatting_plan_json_val, list): # Plan is usually a list of actions
            plan = formatting_plan_json_val
        elif isinstance(formatting_plan_json_val, dict) and "error" in formatting_plan_json_val: # Plan could be an error dict
            plan = formatting_plan_json_val
        else:
            return json.dumps({"error": f"formatting_plan_json must be JSON string, list or error dict, got {type(formatting_plan_json_val)}."})

        # Error checks on the parsed/received data
        if "error" in original_analysis:
            return json.dumps({"error": f"Original document analysis for validation contains an error: {original_analysis['error']}"})
        if "error" in modified_analysis:
            return json.dumps({"error": f"Modified document analysis for validation contains an error: {modified_analysis['error']}"})

        if isinstance(plan, list) and plan and isinstance(plan[0], dict) and "error" in plan[0]:
             return json.dumps({"error": f"Formatting plan (list) for validation indicates an error: {plan[0]['error']}"})
        elif isinstance(plan, dict) and "error" in plan:
             return json.dumps({"error": f"Formatting plan (dict) for validation indicates an error: {plan['error']}"})

    except json.JSONDecodeError as e: # Catches errors from json.loads if inputs were strings
        return json.dumps({"error": f"Invalid JSON string input for {tool_name}: {e}. Received: {tool_input}"})
    except Exception as e: # Catch other unexpected errors during parsing
        return json.dumps({"error": f"Error processing inputs for {tool_name}: {str(e)}. Received: {tool_input}"})

    # Prepare summaries for the prompt (to keep it manageable)
    original_summary = {
        "path": original_analysis.get("document_path", "N/A"),
        "total_elements": original_analysis.get("summary", {}).get("total_elements"),
        "headings": original_analysis.get("summary", {}).get("heading_count"),
        "paragraphs": original_analysis.get("summary", {}).get("paragraph_count"),
        # Potentially add a few sample element details if concise enough
        "sample_elements_before": original_analysis.get("elements", [])[:2]
    }
    modified_summary = {
        "path": modified_analysis.get("document_path", "N/A"),
        "total_elements": modified_analysis.get("summary", {}).get("total_elements"),
        "headings": modified_analysis.get("summary", {}).get("heading_count"),
        "paragraphs": modified_analysis.get("summary", {}).get("paragraph_count"),
        "sample_elements_after": modified_analysis.get("elements", [])[:2]
    }

    prompt = f"""
You are a meticulous document formatting validator. Your task is to assess if a document formatting process was successful.
You will be given:
1. A summary of the document's state BEFORE formatting.
2. A summary of the document's state AFTER formatting.
3. The formatting PLAN that was supposed to be applied.
4. The user's original GOAL.

User Goal: "{user_goal}"

Formatting Plan:
{json.dumps(plan, indent=2)}

Document Analysis BEFORE Formatting (Summary):
{json.dumps(original_summary, indent=2)}

Document Analysis AFTER Formatting (Summary):
{json.dumps(modified_summary, indent=2)}

Based on all this information, provide a validation report in JSON format. The report should include:
- "plan_adherence_summary": (string) Briefly assess if the changes suggested by the plan seem to be reflected in the 'AFTER' summary compared to 'BEFORE'. Highlight any major discrepancies if evident from the summaries.
- "goal_achievement_summary": (string) Qualitatively assess if the user's goal appears to have been met, considering the plan and the changes.
- "overall_assessment": (string, e.g., "Excellent", "Good", "Fair", "Needs Improvement") Your overall qualitative score.
- "suggestions_for_improvement": (list of strings) Specific, actionable suggestions if further improvements can be made or if some plan items were missed. If it looks perfect, say so.

Example of desired JSON output:
{{
  "plan_adherence_summary": "The plan to standardize fonts to Calibri 12pt appears largely successful based on the 'AFTER' summary. Heading styles also seem updated as per plan.",
  "goal_achievement_summary": "The document now appears more professional and consistent, aligning with the user's goal.",
  "overall_assessment": "Good",
  "suggestions_for_improvement": [
    "Consider checking paragraph spacing consistency for list items.",
    "Verify if table formatting also needs to be standardized (if tables exist)."
  ]
}}

IMPORTANT: Only output the JSON object for the report. Do not include any other text, explanations, or markdown formatting.
Focus on what can be inferred from the provided summaries and plan. You don't have the full documents.
"""

    try:
        response = llm.invoke(prompt)
        validation_report_str = response.content

        # Validate if the output is valid JSON
        try:
            json.loads(validation_report_str)
            return validation_report_str
        except json.JSONDecodeError as json_err:
            error_msg = f"LLM returned invalid JSON for validation report. Error: {json_err}. LLM Output: {validation_report_str}"
            print(error_msg)
            return json.dumps({"error": "LLM generated invalid JSON validation report.", "details": error_msg})

    except Exception as e:
        error_msg = f"Error during LLM call in {tool_name}: {e}"
        print(error_msg)
        return json.dumps({"error": error_msg})

if __name__ == '__main__':
    # Create a dummy doc for testing tools
    from docx import Document as PythonDocXDocument # Alias to avoid confusion

    # Create a dummy 'test_document.docx'
    test_doc_content = PythonDocXDocument()
    test_doc_content.add_heading("Main Title", level=1)
    test_doc_content.add_paragraph("This is the first paragraph. It has some initial text.")
    test_doc_content.add_heading("Subtitle A", level=2)
    test_doc_content.add_paragraph("Another paragraph here, with more content to analyze.")
    test_doc_content.save("test_document.docx")
    print("Created 'test_document.docx' for tool testing.")

    print("\n--- Testing analyze_document_structure ---")
    analysis_json_str = analyze_document_structure.invoke({"doc_path": "test_document.docx"})
    print("Raw JSON Analysis Output:")
    print(analysis_json_str)
    try:
        analysis_data = json.loads(analysis_json_str)
        print("\nParsed Analysis Data (Summary):")
        print(f"  Document Path: {analysis_data.get('document_path')}")
        print(f"  Total Elements: {analysis_data.get('summary', {}).get('total_elements')}")
        print(f"  Paragraphs: {analysis_data.get('summary', {}).get('paragraph_count')}")
        print(f"  Headings: {analysis_data.get('summary', {}).get('heading_count')}")
        if "error" in analysis_data:
             print(f"  Error in analysis: {analysis_data['error']}")
        analysis_json_for_plan = analysis_json_str # Keep the JSON string for the next tool
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON from analyze_document_structure: {e}")
        analysis_json_for_plan = json.dumps({"error": "Analysis JSON was invalid."})


    print("\n--- Testing create_formatting_plan ---")
    # Ensure GROQ_API_KEY is available or this test might be skipped/fail if llm is None
    if llm:
        plan_json_str = create_formatting_plan.invoke({
            "document_analysis_json": analysis_json_for_plan,
            "user_goal": "Make it look very professional for a business presentation"
        })
        print("Raw JSON Plan Output:")
        print(plan_json_str)
        try:
            plan_data = json.loads(plan_json_str)
            print("\nParsed Plan Data (first item if exists):")
            if isinstance(plan_data, list) and plan_data:
                print(plan_data[0])
            elif "error" in plan_data:
                 print(f"  Error in plan generation: {plan_data['error']}")
            plan_for_apply = plan_json_str # Keep the JSON string
        except json.JSONDecodeError as e:
            print(f"Error parsing JSON from create_formatting_plan: {e}")
            plan_for_apply = json.dumps([{"action": "error", "details": "Plan JSON was invalid."}])
    else:
        print("Skipping create_formatting_plan test as LLM is not initialized (check API key).")
        # Create a dummy plan string for subsequent tests
        plan_for_apply = json.dumps([
            {"action": "set_font", "scope": "all_body_paragraphs", "font_name": "Arial", "size": 12},
            {"action": "set_heading_style", "level": 1, "font_name": "Arial Black", "size": 18, "bold": True}
        ])
        print(f"Using dummy plan for next steps: {plan_for_apply}")


    print("\n--- Testing apply_contextual_formatting ---")
    # Ensure analysis_json_for_plan is available from the analyze_document_structure step
    if 'analysis_json_for_plan' not in locals():
        # If analyze_document_structure failed or was skipped, create a dummy analysis
        print("Warning: analysis_json_for_plan not found, creating dummy for apply_contextual_formatting test.")
        dummy_analysis_dict = {
            "document_path": "test_document.docx",
            "summary": {"total_elements": 2, "paragraph_count": 1, "heading_count": 1},
            "elements": [
                {"type": "heading", "level": 1, "text": "Main Title", "paragraph_index": 0, "style_name": "Heading 1"},
                {"type": "paragraph", "text": "This is the first paragraph.", "paragraph_index": 1, "style_name": "Normal"}
            ]
        }
        analysis_json_for_plan = json.dumps(dummy_analysis_dict)

    apply_result_json_str = apply_contextual_formatting.invoke({
        "doc_path": "test_document.docx",
        "formatting_plan_json": plan_for_apply,
        "document_analysis_json": analysis_json_for_plan, # New argument
        "output_doc_path": "test_document_modified.docx"
    })
    print("Raw Apply Result JSON Output:")
    print(apply_result_json_str)
    try:
        apply_result_data = json.loads(apply_result_json_str)
        print("\nParsed Apply Result:")
        print(apply_result_data)
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON from apply_contextual_formatting: {e}")


    print("\n--- Testing validate_formatting_result ---")
    # Create dummy analysis strings for original and modified docs for the test
    # In a real agent flow, these would come from actual analyze_document_structure calls
    dummy_original_analysis_data = {
        "document_path": "test_document.docx",
        "summary": {"total_elements": 2, "paragraph_count": 1, "heading_count": 1, "font_predominant": "Times New Roman"},
        "elements": [
            {"type": "heading", "level": 1, "text": "Old Title", "paragraph_index": 0, "runs": [{"font_name": "Times New Roman", "font_size": 16}]},
            {"type": "paragraph", "text": "Old text.", "paragraph_index": 1, "runs": [{"font_name": "Times New Roman", "font_size": 12}]}
        ]
    }
    dummy_modified_analysis_data = {
        "document_path": "test_document_modified.docx",
        "summary": {"total_elements": 2, "paragraph_count": 1, "heading_count": 1, "font_predominant": "Arial"},
         "elements": [
            {"type": "heading", "level": 1, "text": "Main Title", "paragraph_index": 0, "runs": [{"font_name": "Arial", "font_size": 18, "bold": True}]},
            {"type": "paragraph", "text": "This is the first paragraph.", "paragraph_index": 1, "runs": [{"font_name": "Arial", "font_size": 12}]}
        ]
    }
    original_doc_analysis_json_for_validation = json.dumps(dummy_original_analysis_data)
    modified_doc_analysis_json_for_validation = json.dumps(dummy_modified_analysis_data)
    user_goal_for_validation = "Make it look very professional for a business presentation using Arial font."

    if llm: # Only run if LLM is available
        validation_json_str = validate_formatting_result.invoke({
            "original_doc_analysis_json": original_doc_analysis_json_for_validation,
            "modified_doc_analysis_json": modified_doc_analysis_json_for_validation,
            "formatting_plan_json": plan_for_apply, # This is the plan generated earlier
            "user_goal": user_goal_for_validation
        })
        print("Raw Validation Report JSON Output:")
        print(validation_json_str)
        try:
            validation_data = json.loads(validation_json_str)
            print("\nParsed Validation Report:")
            print(f"  Overall Assessment: {validation_data.get('overall_assessment')}")
            print(f"  Goal Achievement: {validation_data.get('goal_achievement_summary')}")
            if "error" in validation_data:
                print(f"  Error in validation: {validation_data['error']}")
        except json.JSONDecodeError as e:
            print(f"Error parsing JSON from validate_formatting_result: {e}")
    else:
        print("Skipping validate_formatting_result test as LLM is not initialized.")
        validation_json_str = json.dumps({"error": "LLM not available for validation test."}) # Placeholder
        print(validation_json_str)


    # Clean up dummy files
    import os
    os.remove("test_document.docx")
    if os.path.exists("test_document_modified.docx"):
        os.remove("test_document_modified.docx")
    print("\nCleaned up test files.")
