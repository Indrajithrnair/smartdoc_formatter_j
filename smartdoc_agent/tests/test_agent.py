import unittest
import os
from unittest.mock import patch, MagicMock

# Attempt to import agent and config, handling potential import errors
try:
    from smartdoc_agent.core.agent import DocumentFormattingAgent
    from smartdoc_agent.config import get_groq_api_key
except ImportError:
    import sys
    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
    from core.agent import DocumentFormattingAgent
    from config import get_groq_api_key

# To create a dummy doc for tests
from docx import Document as PythonDocXDocument


class TestDocumentFormattingAgent(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        cls.sample_doc_path = "temp_agent_test_doc.docx"
        cls.sample_output_path = "temp_agent_test_doc_formatted.docx"

        # Create a dummy document
        doc = PythonDocXDocument()
        doc.add_heading("Agent Test", level=1)
        doc.add_paragraph("A paragraph for the agent to process.")
        doc.save(cls.sample_doc_path)

        # Ensure there's a .env file or the key is in environment for tests
        # For CI environments, GROQ_API_KEY should be set as an env variable
        if not get_groq_api_key():
            # Create a dummy .env if not present and no env var
            # This is mainly for local testing convenience
            if not os.path.exists(os.path.join(os.path.dirname(__file__), '..', '.env')):
                env_path = os.path.join(os.path.dirname(__file__), '..', '.env')
                with open(env_path, 'w') as f:
                    f.write('GROQ_API_KEY="DUMMY_KEY_FOR_TEST_INITIALIZATION"\n')
                print(f"Created dummy .env at {env_path} for test initialization. Agent might fail if DUMMY_KEY is not valid and actual LLM call is made.")


    @classmethod
    def tearDownClass(cls):
        if os.path.exists(cls.sample_doc_path):
            os.remove(cls.sample_doc_path)
        if os.path.exists(cls.sample_output_path):
            os.remove(cls.sample_output_path)

        # Clean up dummy .env if it was created by this test
        dummy_env_path = os.path.join(os.path.dirname(__file__), '..', '.env')
        # A bit risky if there was a REAL .env with this dummy key.
        # A better approach would be to mock config.get_groq_api_key()
        # if os.path.exists(dummy_env_path):
        #     with open(dummy_env_path, 'r') as f:
        #         content = f.read()
        #     if 'GROQ_API_KEY="DUMMY_KEY_FOR_TEST_INITIALIZATION"' in content:
        #         # os.remove(dummy_env_path) # Commented out for safety for now
        #         pass


    @patch('smartdoc_agent.core.agent.ChatGroq') # Mock the LLM
    @patch('smartdoc_agent.core.agent.AgentExecutor') # Mock the AgentExecutor
    def test_agent_initialization_and_run_structure(self, MockAgentExecutor, MockChatGroq):
        """Test agent initializes and its run method can be called, mocking external dependencies."""

        # Setup mocks
        mock_llm_instance = MockChatGroq.return_value
        mock_agent_executor_instance = MockAgentExecutor.return_value
        mock_agent_executor_instance.invoke.return_value = {"output": "Mocked agent response"}

        try:
            agent = DocumentFormattingAgent()
            self.assertIsNotNone(agent.llm, "LLM should be initialized.")
            self.assertIsNotNone(agent.agent_executor, "AgentExecutor should be initialized.")

            # Test the run method
            user_query = "Make this document professional."
            result = agent.run(user_query, self.sample_doc_path, self.sample_output_path)

            # Check if invoke was called with expected structure
            mock_agent_executor_instance.invoke.assert_called_once()
            call_args = mock_agent_executor_instance.invoke.call_args[0][0]
            self.assertIn("input", call_args)
            self.assertIn(user_query, call_args["input"])
            self.assertIn(self.sample_doc_path, call_args["input"])

            self.assertEqual(result, "Mocked agent response")

        except ValueError as e:
            # This might happen if GROQ_API_KEY is not set and the mock isn't perfect
            # or if the dummy key is actually used by ChatGroq initialization logic.
            if "Groq API key not found" in str(e) or "API key" in str(e).lower():
                self.skipTest(f"GROQ_API_KEY related error during setup: {e}. Ensure key is available or mock is complete.")
            else:
                raise # Re-raise other ValueErrors

    @patch('smartdoc_agent.config.get_groq_api_key', return_value=None)
    def test_agent_initialization_fails_without_api_key(self, mock_get_key):
        """Test that agent initialization raises ValueError if API key is missing."""
        with self.assertRaisesRegex(ValueError, "Groq API key not found"):
            DocumentFormattingAgent()


    # More tests will be added here as the agent's capabilities grow.
    # For example, testing specific tool selection logic, prompt effectiveness, etc.
    # These would require more sophisticated mocking or integration-style tests.

if __name__ == '__main__':
    unittest.main()
