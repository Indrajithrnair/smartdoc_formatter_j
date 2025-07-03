import unittest
import os
from unittest.mock import patch, MagicMock # Correctly import patch and MagicMock
from docx import Document as PythonDocXDocument # Alias to avoid confusion

# Attempt to import tools, handling potential import errors if run standalone
try:
    from smartdoc_agent.core.tools import (
        analyze_document_structure,
        create_formatting_plan,
        apply_contextual_formatting,
        validate_formatting_result
    )
    from smartdoc_agent.utils.document_utils import load_document, save_document
except ImportError:
    # This is to allow running the test file directly for development,
    # assuming it's in the tests/ directory and smartdoc_agent is in the parent.
    import sys
    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
    from core.tools import (
        analyze_document_structure,
        create_formatting_plan,
        apply_contextual_formatting,
        validate_formatting_result
    )
    from utils.document_utils import load_document, save_document


class TestFormattingTools(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        """Set up for all tests; create a sample document."""
        cls.test_doc_path = "temp_test_document_for_tools.docx"
        cls.test_output_doc_path = "temp_test_document_for_tools_modified.docx"

        doc = PythonDocXDocument()
        doc.add_heading("Test Document Title", level=1)
        doc.add_paragraph("This is a paragraph for testing.")
        doc.add_heading("Another Heading", level=2)
        doc.add_paragraph("Second paragraph with some more text.")
        save_document(doc, cls.test_doc_path)

    @classmethod
    def tearDownClass(cls):
        """Clean up after all tests; remove the sample document."""
        if os.path.exists(cls.test_doc_path):
            os.remove(cls.test_doc_path)
        if os.path.exists(cls.test_output_doc_path):
            os.remove(cls.test_output_doc_path)

    def test_analyze_document_structure_runs(self):
        """Test that analyze_document_structure runs and returns a JSON string."""
        import json
        result_json_str = analyze_document_structure.invoke({"doc_path": self.test_doc_path})
        self.assertIsInstance(result_json_str, str)

        try:
            result_data = json.loads(result_json_str)
        except json.JSONDecodeError:
            self.fail("analyze_document_structure did not return valid JSON.")

        self.assertIsInstance(result_data, dict)
        self.assertEqual(result_data.get("document_path"), self.test_doc_path)
        self.assertIn("elements", result_data)
        self.assertIn("summary", result_data)
        self.assertTrue(len(result_data["elements"]) > 0)
        self.assertTrue(result_data["summary"]["total_elements"] > 0)
        # Check if one of the known texts from setUpClass is present in the analysis
        found_title = any(el["text"] == "Test Document Title" for el in result_data["elements"])
        self.assertTrue(found_title, "Expected content not found in analysis elements.")


    @patch('smartdoc_agent.core.tools.llm') # Mock the LLM used by the tool
    def test_create_formatting_plan_runs(self, mock_llm):
        """Test that create_formatting_plan runs and returns a JSON string plan."""
        import json

        # Mock the LLM's response
        mock_llm_response_content = json.dumps([
            {"action": "set_font", "scope": "all_body_paragraphs", "font_name": "Arial", "size": 12},
            {"action": "set_heading_style", "level": 1, "font_name": "Arial Black", "size": 16, "bold": True}
        ])
        mock_llm.invoke.return_value = MagicMock(content=mock_llm_response_content)

        # Sample analysis (must be a JSON string)
        sample_analysis_dict = {
            "document_path": self.test_doc_path,
            "summary": {"total_elements": 4, "paragraph_count": 2, "heading_count": 2},
            "elements": [
                {"type": "heading", "level": 1, "text": "Test Document Title", "style_name": "Heading 1", "runs": []},
                {"type": "paragraph", "text": "This is a paragraph for testing.", "style_name": "Normal", "runs": []}
            ]
        }
        sample_analysis_json = json.dumps(sample_analysis_dict)
        user_goal = "Make it look professional."

        result_plan_json = create_formatting_plan.invoke({
            "document_analysis_json": sample_analysis_json,
            "user_goal": user_goal
        })

        self.assertIsInstance(result_plan_json, str)
        try:
            plan_data = json.loads(result_plan_json)
        except json.JSONDecodeError:
            self.fail("create_formatting_plan did not return valid JSON.")

        self.assertIsInstance(plan_data, list)
        self.assertTrue(len(plan_data) > 0)
        self.assertIn("action", plan_data[0])
        mock_llm.invoke.assert_called_once() # Check that LLM was called

    def test_apply_contextual_formatting_runs(self):
        """Test that apply_contextual_formatting runs and returns a string."""
        import json
        # This is a basic test for the stub.
        # The plan should now be a JSON string.
        sample_plan_list = [
            {"action": "set_font", "scope": "all_body_paragraphs", "font_name": "Times New Roman", "size": 12},
            {"action": "set_heading_style", "level": 1, "font_name": "TNR Bold", "size": 14, "bold": True}
        ]
        sample_plan_json = json.dumps(sample_plan_list)

        # Sample analysis JSON string (needed by the refined tool)
        sample_analysis_dict = {
            "document_path": self.test_doc_path,
            "summary": {"total_elements": 1, "paragraph_count": 1, "heading_count": 0},
            "elements": [{"type": "paragraph", "text": "Text", "paragraph_index": 0}]
        }
        sample_analysis_json = json.dumps(sample_analysis_dict)

        # Mock the actual document utility functions that perform changes
        with patch('smartdoc_agent.core.tools.load_document') as mock_load, \
             patch('smartdoc_agent.core.tools.save_document') as mock_save, \
             patch('smartdoc_agent.core.tools.apply_set_font_action') as mock_apply_font, \
             patch('smartdoc_agent.core.tools.apply_set_heading_style_action') as mock_apply_heading:

            mock_doc_instance = MagicMock(spec=PythonDocXDocument)
            mock_load.return_value = mock_doc_instance

            result_json_str = apply_contextual_formatting.invoke({
                "doc_path": self.test_doc_path,
                "formatting_plan_json": sample_plan_json,
                "document_analysis_json": sample_analysis_json, # New argument
                "output_doc_path": self.test_output_doc_path
            })
            self.assertIsInstance(result_json_str, str)

            try:
                result_data = json.loads(result_json_str)
            except json.JSONDecodeError:
                self.fail("apply_contextual_formatting did not return valid JSON.")

            self.assertEqual(result_data.get("status"), "success")
            self.assertTrue(os.path.exists(self.test_output_doc_path)) # Still creates the file due to stub

            # Check if our action handlers were called
            mock_load.assert_called_once_with(self.test_doc_path)
            mock_apply_font.assert_called_once()
            mock_apply_heading.assert_called_once()
            mock_save.assert_called_once_with(mock_doc_instance, self.test_output_doc_path)


    def test_validate_formatting_result_runs(self):
        """Test that validate_formatting_result runs and returns a string."""
        # This is a basic test for the stub.
        # Ensure the modified document exists for validation
        if not os.path.exists(self.test_output_doc_path):
             # Create a dummy one if apply_contextual_formatting hasn't created it yet
            doc = load_document(self.test_doc_path)
            save_document(doc, self.test_output_doc_path)

        sample_plan = "1. Change font. 2. Check headings."
        user_goal = "Check professionalism."
        result = validate_formatting_result.invoke({
            "original_doc_path": self.test_doc_path,
            "modified_doc_path": self.test_output_doc_path,
            "formatting_plan": sample_plan,
            "user_goal": user_goal
        })
        self.assertIsInstance(result, str)
        self.assertIn(f"Validation Report for goal '{user_goal}'", result)

if __name__ == '__main__':
    unittest.main()
