from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH # Already here, good.

# Ensure it's available for the functions in this file too.
# It seems it was already imported, so the issue might be elsewhere or this is just a check.

def load_document(file_path: str) -> Document:
    """Loads a .docx document from the given file path."""
    try:
        document = Document(file_path)
        return document
    except Exception as e:
        print(f"Error loading document {file_path}: {e}")
        raise

def save_document(document: Document, file_path: str) -> None:
    """Saves the given Document object to a .docx file."""
    try:
        document.save(file_path)
        print(f"Document saved to {file_path}")
    except Exception as e:
        print(f"Error saving document to {file_path}: {e}")
        raise

def extract_text_from_paragraphs(document: Document) -> list[str]:
    """Extracts text from all paragraphs in the document."""
    texts = []
    for para in document.paragraphs:
        texts.append(para.text)
    return texts

def extract_headings(document: Document) -> list[dict[str, str]]:
    """
    Extracts headings and their levels from the document.
    Assumes headings are styled with Word's default heading styles (e.g., "Heading 1", "Heading 2").
    """
    headings = []
    for para in document.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.split(" ")[-1])
            except ValueError:
                level = 0 # Or handle as an unknown heading level
            headings.append({"level": level, "text": para.text})
    return headings

# Example of a more complex utility function placeholder for later
# def change_paragraph_font(paragraph, font_name: str, font_size: Pt):
#     for run in paragraph.runs:
#         run.font.name = font_name
#         run.font.size = Pt(font_size)

if __name__ == '__main__':
    # This is for basic testing of the utility functions.
    # You would need a sample .docx file named 'sample.docx' in the same directory.
    # Create a dummy document for testing if 'sample.docx' doesn't exist
    doc = Document()
    doc.add_heading('Test Document', level=0)
    doc.add_paragraph('This is a test paragraph.')
    doc.add_heading('Heading 1', level=1)
    doc.add_paragraph('Another paragraph under Heading 1.')
    doc.add_heading('Heading 2', level=2)
    doc.add_paragraph('Paragraph under Heading 2.')
    save_document(doc, 'sample_test_doc.docx')

    print("Testing with 'sample_test_doc.docx'")
    loaded_doc = load_document('sample_test_doc.docx')

    print("\nExtracted Paragraph Texts:")
    texts = extract_text_from_paragraphs(loaded_doc)
    for text in texts:
        print(f"- {text}")

    print("\nExtracted Headings:")
    headings_info = extract_headings(loaded_doc)
    for heading in headings_info:
        print(f"- Level {heading['level']}: {heading['text']}")

    print("\nDetailed Document Analysis:")
    detailed_analysis_data = get_document_analysis(loaded_doc) # Use the loaded_doc directly
    for item in detailed_analysis_data["elements"]:
        print(f"- Type: {item['type']}")
        if item['type'] == 'heading':
            print(f"  Level: {item['level']}")
        print(f"  Text: {item['text'][:50]}...") # Print first 50 chars
        if 'style' in item:
            print(f"  Style: {item['style']}")
        if 'runs' in item:
            # print(f"  Runs: {item['runs']}") # This can be verbose
            if item['runs']:
                first_run = item['runs'][0]
                print(f"  First run font: {first_run.get('font_name', 'N/A')}, Size: {first_run.get('font_size', 'N/A')}, Bold: {first_run.get('bold', 'N/A')}")

    # Clean up the dummy file
    import os
    os.remove('sample_test_doc.docx')
    print("\nCleaned up sample_test_doc.docx")

def get_run_details(run) -> dict:
    """Extracts details from a run."""
    return {
        "text": run.text,
        "font_name": run.font.name,
        "font_size": run.font.size.pt if run.font.size else None, # Size in points
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
    }

def get_paragraph_details(para, para_index: int) -> dict:
    """Extracts details from a paragraph, including its runs."""
    alignment_name = None
    if para.alignment is not None:
        # para.alignment should be a member of WD_ALIGN_PARAGRAPH enum
        # Accessing .name gives the string representation like 'LEFT', 'CENTER'
        try:
            alignment_name = para.alignment.name
        except AttributeError: # Should not happen if para.alignment is a valid enum member
            alignment_name = str(para.alignment) # Fallback to string of the value

    para_info = {
        "paragraph_index": para_index,
        "text": para.text,
        "style_name": para.style.name if para.style else "Default Paragraph Font",
        "alignment": alignment_name,
        "runs": [get_run_details(run) for run in para.runs]
    }
    # Heading level detection
    if para.style and para.style.name.startswith('Heading'):
        try:
            level = int(para.style.name.split(' ')[-1])
            para_info["type"] = "heading"
            para_info["level"] = level
        except ValueError:
            para_info["type"] = "paragraph" # Could be a custom heading style not ending in a number
            para_info["level"] = 0
    elif para.style and para.style.name == "Title": # Common style for document title
        para_info["type"] = "heading"
        para_info["level"] = 0
    else:
        para_info["type"] = "paragraph"

    return para_info

def get_document_analysis(document: Document) -> dict:
    """
    Analyzes a Document object and extracts detailed information about its elements.
    Returns a dictionary with a list of element details.
    """
    analysis = {"elements": []}
    for i, para in enumerate(document.paragraphs):
        para_details = get_paragraph_details(para, i)
        analysis["elements"].append(para_details)

    # Future: Add analysis for tables, lists, images, sections, etc.
    return analysis

# --- Formatting Application Utilities ---

def set_paragraph_font_properties(paragraph, font_name: str = None, size_pt: float = None, bold: bool = None, italic: bool = None, underline: bool = None):
    """Applies font properties to all runs in a paragraph."""
    for run in paragraph.runs:
        if font_name:
            run.font.name = font_name
        if size_pt:
            run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if underline is not None:
            run.underline = underline

def set_paragraph_spacing_properties(paragraph, spacing_before_pt: float = None, spacing_after_pt: float = None, line_spacing_rule: float = None): # line_spacing_rule e.g. 1.0, 1.5, 2.0
    """Applies spacing properties to a paragraph."""
    if spacing_before_pt is not None:
        paragraph.paragraph_format.space_before = Pt(spacing_before_pt)
    if spacing_after_pt is not None:
        paragraph.paragraph_format.space_after = Pt(spacing_after_pt)
    if line_spacing_rule is not None:
        paragraph.paragraph_format.line_spacing = line_spacing_rule


def set_paragraph_alignment_properties(paragraph, alignment: str = None): # alignment: "LEFT", "CENTER", "RIGHT", "JUSTIFY"
    """Applies alignment to a paragraph."""
    if alignment:
        try:
            align_enum = getattr(WD_ALIGN_PARAGRAPH, alignment.upper(), None)
            if align_enum is not None:
                paragraph.alignment = align_enum
            else:
                print(f"Warning: Invalid alignment value '{alignment}'. Skipping.")
        except Exception as e:
            print(f"Warning: Exception setting alignment '{alignment}': {e}")


# More specific action handlers to be called by the tool:

def apply_set_font_action(doc: Document, elements_details: list, action: dict):
    """
    Applies font settings based on the action dictionary.
    Action: {"action": "set_font", "scope": "all_paragraphs" | "headings_level_X" | "paragraph_index_N",
             "font_name": "Arial", "size": 12, "bold": false, "italic": false}
    """
    print(f"Applying font action: {action}")
    scope = action.get("scope")
    font_name = action.get("font_name")
    size = action.get("size")
    bold = action.get("bold")
    italic = action.get("italic")
    underline = action.get("underline")

    target_paras = []
    if scope == "all_paragraphs":
        target_paras = doc.paragraphs
    elif scope and scope.startswith("headings_level_"):
        try:
            level = int(scope.split("_")[-1])
            # Iterate through elements_details to find matching paragraphs in the doc
            for i, el_detail in enumerate(elements_details):
                if el_detail.get("type") == "heading" and el_detail.get("level") == level:
                    # Ensure paragraph index is valid
                    if el_detail["paragraph_index"] < len(doc.paragraphs):
                        target_paras.append(doc.paragraphs[el_detail["paragraph_index"]])
        except ValueError:
            print(f"Warning: Invalid heading level in scope '{scope}'.")
    elif scope and scope.startswith("paragraph_index_"):
        try:
            idx = int(scope.split("_")[-1])
            if 0 <= idx < len(doc.paragraphs):
                target_paras.append(doc.paragraphs[idx])
            else:
                print(f"Warning: Paragraph index {idx} out of bounds.")
        except ValueError:
            print(f"Warning: Invalid paragraph index in scope '{scope}'.")
    elif scope == "all_body_paragraphs": # Distinguish from headings
         for i, el_detail in enumerate(elements_details):
            if el_detail.get("type") == "paragraph": # Not a heading
                if el_detail["paragraph_index"] < len(doc.paragraphs):
                    target_paras.append(doc.paragraphs[el_detail["paragraph_index"]])
    else:
        print(f"Warning: Unknown or unsupported scope '{scope}' for set_font.")
        return

    for para in target_paras:
        set_paragraph_font_properties(para, font_name, size, bold, italic, underline)
    print(f"Applied font to {len(target_paras)} paragraphs for scope '{scope}'.")


def apply_set_heading_style_action(doc: Document, elements_details: list, action: dict):
    """
    Applies style (font, size, bold etc.) to headings of a specific level.
    Action: {"action": "set_heading_style", "level": 1, "font_name": "Calibri Light",
             "size": 18, "bold": true, "spacing_after": 12}
    """
    print(f"Applying heading style action: {action}")
    level = action.get("level")
    font_name = action.get("font_name")
    size = action.get("size")
    bold = action.get("bold")
    italic = action.get("italic")
    underline = action.get("underline")
    spacing_after_pt = action.get("spacing_after")
    # keep_with_next = action.get("keep_with_next") # TODO

    target_paras = []
    # Iterate through elements_details to find matching paragraphs in the doc
    for i, el_detail in enumerate(elements_details):
        if el_detail.get("type") == "heading" and el_detail.get("level") == level:
            if el_detail["paragraph_index"] < len(doc.paragraphs):
                target_paras.append(doc.paragraphs[el_detail["paragraph_index"]])

    for para in target_paras:
        set_paragraph_font_properties(para, font_name, size, bold, italic, underline)
        if spacing_after_pt is not None:
             set_paragraph_spacing_properties(para, spacing_after_pt=spacing_after_pt)
        # if keep_with_next is not None:
        #     para.paragraph_format.keep_with_next = keep_with_next
    print(f"Applied style to {len(target_paras)} Level {level} headings.")


def apply_set_paragraph_spacing_action(doc: Document, elements_details: list, action: dict):
    """
    Applies paragraph spacing settings.
    Action: {"action": "set_paragraph_spacing", "scope": "all_paragraphs",
             "spacing_before": 0, "spacing_after": 6, "line_spacing": 1.15}
    """
    print(f"Applying paragraph spacing action: {action}")
    scope = action.get("scope")
    spacing_before = action.get("spacing_before")
    spacing_after = action.get("spacing_after")
    line_spacing = action.get("line_spacing")

    target_paras = []
    if scope == "all_paragraphs":
        target_paras = doc.paragraphs
    elif scope == "all_body_paragraphs":
         for i, el_detail in enumerate(elements_details):
            if el_detail.get("type") == "paragraph": # Not a heading
                if el_detail["paragraph_index"] < len(doc.paragraphs):
                    target_paras.append(doc.paragraphs[el_detail["paragraph_index"]])
    # Add more scopes if needed (e.g., specific paragraph indices, styles)
    else:
        print(f"Warning: Unknown or unsupported scope '{scope}' for set_paragraph_spacing.")
        return

    for para in target_paras:
        set_paragraph_spacing_properties(para, spacing_before, spacing_after, line_spacing)
    print(f"Applied spacing to {len(target_paras)} paragraphs for scope '{scope}'.")


def apply_set_alignment_action(doc: Document, elements_details: list, action: dict):
    """
    Applies text alignment.
    Action: {"action": "set_alignment", "scope": "headings_level_1", "alignment": "LEFT"}
    """
    print(f"Applying alignment action: {action}")
    scope = action.get("scope")
    alignment = action.get("alignment")

    target_paras = []
    if scope == "all_paragraphs":
        target_paras = doc.paragraphs
    elif scope and scope.startswith("headings_level_"):
        try:
            level = int(scope.split("_")[-1])
            for i, el_detail in enumerate(elements_details):
                if el_detail.get("type") == "heading" and el_detail.get("level") == level:
                     if el_detail["paragraph_index"] < len(doc.paragraphs):
                        target_paras.append(doc.paragraphs[el_detail["paragraph_index"]])
        except ValueError:
            print(f"Warning: Invalid heading level in scope '{scope}'.")
    elif scope == "all_body_paragraphs":
         for i, el_detail in enumerate(elements_details):
            if el_detail.get("type") == "paragraph": # Not a heading
                if el_detail["paragraph_index"] < len(doc.paragraphs):
                    target_paras.append(doc.paragraphs[el_detail["paragraph_index"]])
    else:
        print(f"Warning: Unknown or unsupported scope '{scope}' for set_alignment.")
        return

    for para in target_paras:
        set_paragraph_alignment_properties(para, alignment)
    print(f"Applied alignment to {len(target_paras)} paragraphs for scope '{scope}'.")

# TODO: Implement other action handlers like:
# def apply_ensure_consistent_style_action(doc: Document, action: dict): ...
# def apply_theme_action(doc: Document, action: dict): ...

def apply_fix_font_inconsistencies_action(doc: Document, elements_details: list, action: dict):
    """
    Attempts to unify fonts across the document based on a target font and size.
    Action: {"action": "fix_font_inconsistencies", "target_font_name": "Calibri", "target_font_size": 11}
    This is a broad approach. More nuanced logic could be added to preserve specific formatting
    (e.g., for code blocks, or manually emphasized text if distinguishable).
    """
    print(f"Applying font inconsistency fix: {action}")
    target_font_name = action.get("target_font_name")
    target_font_size_pt = action.get("target_font_size") # Assuming this is in Pt

    if not target_font_name and not target_font_size_pt:
        print("Warning: No target font name or size provided for fix_font_inconsistencies. Skipping.")
        return

    changed_elements_count = 0
    for para_idx, para_detail in enumerate(elements_details):
        # We operate on the doc.paragraphs directly using original indices
        if para_idx >= len(doc.paragraphs):
            continue

        paragraph = doc.paragraphs[para_idx]

        # Skip headings if plan also has set_heading_style, to avoid conflicts or apply selectively.
        # For a generic "fix inconsistencies", we might apply to all, or let LLM be more specific with scope.
        # Current plan from LLM for "make headings bold" also included this, so it might apply to headings too.
        # Let's assume for now it applies to all runs in all paragraphs unless scope is narrowed by LLM.

        for run in paragraph.runs:
            applied_change_to_run = False
            if target_font_name and run.font.name != target_font_name:
                # print(f"  Run '{run.text[:20]}' changing font from {run.font.name} to {target_font_name}")
                run.font.name = target_font_name
                applied_change_to_run = True
            if target_font_size_pt and run.font.size != Pt(target_font_size_pt):
                # print(f"  Run '{run.text[:20]}' changing size from {run.font.size} to {Pt(target_font_size_pt)}")
                run.font.size = Pt(target_font_size_pt)
                applied_change_to_run = True
            if applied_change_to_run:
                changed_elements_count +=1 # Count runs changed, not paragraphs

    print(f"Applied font inconsistency fix to {changed_elements_count} runs.")
